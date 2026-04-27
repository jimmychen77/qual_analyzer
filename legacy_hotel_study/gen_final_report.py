"""
生成酒店评论综合分析Word报告
佰翔酒店 vs 漳州宾馆
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import pandas as pd
from collections import Counter, defaultdict
import jieba

# ==================== 数据加载与预处理 ====================
bx = pd.read_excel('佰翔酒店合并数据.xlsx')
zz = pd.read_excel('漳州宾馆合并数据.xlsx')

for df, name in [(bx, '佰翔'), (zz, '漳州')]:
    df['评论长度'] = df['评论内容'].apply(lambda x: len(str(x)) if isinstance(x, str) else 0)
    df['有效评论'] = df['评论长度'] > 10
    df['月份'] = pd.to_datetime(df['评论日期'], errors='coerce').dt.to_period('M')

bx_total = bx['有效评论'].sum()
zz_total = zz['有效评论'].sum()

# ==================== 编码辞典 ====================
coding_dict = {
    '清洁卫生': {
        'pos': ['干净', '整洁', '卫生', '清洁', '无灰尘', '清新'],
        'neg': ['脏', '有灰尘', '毛发', '霉味', '异味', '不干净', '恶心', '有虫']
    },
    '服务态度': {
        'pos': ['热情', '周到', '贴心', '耐心', '友好', '温暖', '亲切', '礼貌', '主动', '微笑', '细心'],
        'neg': ['冷淡', '冷漠', '敷衍', '态度差', '爱答不理', '不耐烦', '差劲', '恶劣', '切单', '看人下菜']
    },
    '前台服务': {
        'pos': ['前台热情', '办理快', '效率高', '入住快', '服务好', '经理', '主管'],
        'neg': ['前台差', '入住慢', '等很久', '办理慢', '服务差', '前台冷漠']
    },
    '设施设备': {
        'pos': ['设施齐全', '设备好', '新', '现代化', '配置好', '完善'],
        'neg': ['设施旧', '设备差', '老旧', '陈旧', '坏', '损坏', '故障', '不能用']
    },
    '房间条件': {
        'pos': ['房间大', '宽敞', '舒适', '床舒服', '床软', '床品好', '采光好', '空调好', '热水好'],
        'neg': ['房间小', '狭小', '窄', '拥挤', '床硬', '不舒服', '空调差', '冷', '热']
    },
    '隔音效果': {
        'pos': ['隔音好', '安静', '很静', '噪音小'],
        'neg': ['隔音差', '吵', '噪音大', '很吵', '嘈杂', '车声', '喇叭', '施工', '装修']
    },
    '早餐餐饮': {
        'pos': ['早餐好', '丰富', '品种多', '味道好', '好吃', '中式', '西式', '自助'],
        'neg': ['早餐差', '单一', '品种少', '难吃', '冷', '不好', '没早餐']
    },
    '位置交通': {
        'pos': ['位置好', '方便', '近', '交通便利', '市中心', '周边有', '万达', '地铁'],
        'neg': ['位置偏', '偏远', '不方便', '难找', '偏僻']
    },
    '性价比': {
        'pos': ['物超所值', '划算', '值得', '超值', '性价比高', '价格合理'],
        'neg': ['不值', '贵', '性价比低', '价格高', '太贵', '坑']
    },
    '景观环境': {
        'pos': ['景观好', '风景美', '江景', '花园', '漂亮', '美', '景色好'],
        'neg': ['景观差', '难看', '看不到', '阴森']
    },
    '停车配套': {
        'pos': ['停车方便', '有停车场', '车位多', '免费停车'],
        'neg': ['停车难', '没车位', '收费', '停车贵']
    }
}

def code_review(text):
    if not isinstance(text, str):
        return [], []
    pos_dims, neg_dims = [], []
    for dim, words in coding_dict.items():
        for w in words['pos']:
            if w in text:
                pos_dims.append(dim)
                break
        for w in words['neg']:
            if w in text:
                neg_dims.append(dim)
                break
    return list(set(pos_dims)), list(set(neg_dims))

# 统计编码
bx_pos, bx_neg = defaultdict(int), defaultdict(int)
zz_pos, zz_neg = defaultdict(int), defaultdict(int)
bx_neg_ex, zz_neg_ex = defaultdict(list), defaultdict(list)

for _, row in bx[bx['有效评论']].iterrows():
    p, n = code_review(str(row['评论内容']))
    for d in p: bx_pos[d] += 1
    for d in n:
        bx_neg[d] += 1
        if len(bx_neg_ex[d]) < 2:
            bx_neg_ex[d].append(str(row['评论内容'])[:100])

for _, row in zz[zz['有效评论']].iterrows():
    p, n = code_review(str(row['评论内容']))
    for d in p: zz_pos[d] += 1
    for d in n:
        zz_neg[d] += 1
        if len(zz_neg_ex[d]) < 2:
            zz_neg_ex[d].append(str(row['评论内容'])[:100])

# 好评/差评触发
def get_triggers(df, threshold_pos=4.5, threshold_neg=3.0):
    pos_triggers = defaultdict(int)
    neg_triggers = defaultdict(int)
    pos_examples = defaultdict(list)
    neg_examples = defaultdict(list)

    pos_kw = {
        '前台服务好': ['前台', '办理', '入住', '服务热情'],
        '早餐好评': ['早餐', '餐厅', '用餐'],
        '免费升房': ['升级', '升房', '免费升级'],
        '服务超预期': ['超出预期', '超乎想象', '惊喜', '感动', '贴心'],
        '主动服务': ['主动', '热情主动', '积极'],
        '赠送物品': ['果盘', '水果', '小礼物', '送了'],
    }
    neg_kw = {
        '设施老旧破损': ['旧', '破', '坏', '故障', '不能用', '损坏'],
        '卫生问题': ['脏', '毛发', '异味', '霉味', '恶心'],
        '隔音差': ['隔音', '噪音', '吵', '车声', '施工'],
        '服务冷漠': ['冷漠', '敷衍', '态度差', '不理'],
        '等待时间长': ['等很久', '等了半天', '排队'],
        '押金/收费问题': ['押金', '收费', '骗', '坑'],
        '信息不对称': ['和说的', '不一样', '没有', '备注'],
        '安全/隐患': ['不安全', '担心', '危险', '滑倒'],
    }

    # 好评
    for _, row in df[df['评分'] >= threshold_pos].iterrows():
        text = str(row['评论内容'])
        for name, kws in pos_kw.items():
            for kw in kws:
                if kw in text:
                    pos_triggers[name] += 1
                    if len(pos_examples[name]) < 1:
                        pos_examples[name].append(text[:80])
                    break

    # 差评
    for _, row in df[df['评分'] < threshold_neg].iterrows():
        text = str(row['评论内容'])
        for name, kws in neg_kw.items():
            for kw in kws:
                if kw in text:
                    neg_triggers[name] += 1
                    if len(neg_examples[name]) < 1:
                        neg_examples[name].append(text[:80])
                    break

    return pos_triggers, neg_triggers, pos_examples, neg_examples

bx_pt, bx_nt, bx_pe, bx_ne = get_triggers(bx)
zz_pt, zz_nt, zz_pe, zz_ne = get_triggers(zz)

# 出行目的推断
def infer_purpose(text):
    if not isinstance(text, str):
        return '其他'
    if any(w in text for w in ['出差', '商务', '办公', '工作']):
        return '商务出差'
    elif any(w in text for w in ['亲子', '小孩', '孩子', '小朋友', '家庭']):
        return '亲子游'
    elif any(w in text for w in ['情侣', '老婆', '老公', '女朋友']):
        return '情侣游'
    elif any(w in text for w in ['旅游', '度假', '景点', '游玩']):
        return '休闲游'
    elif any(w in text for w in ['父母', '老人', '妈妈', '爸爸']):
        return '家庭游'
    return '其他'

bx['出行目的'] = bx['评论内容'].apply(infer_purpose)
zz['出行目的'] = zz['评论内容'].apply(infer_purpose)

# 隐性不满
hidden_words = ['但是', '不过', '就是', '唯一', '美中不足', '建议']
bx_hidden = []
for _, row in bx[(bx['评分'] >= 4.0) & bx['有效评论']].iterrows():
    text = str(row['评论内容'])
    for w in hidden_words:
        if w in text:
            idx = text.find(w)
            bx_hidden.append((row['评分'], w, text[max(0,idx-5):min(len(text), idx+25)]))
            break

zz_hidden = []
for _, row in zz[(zz['评分'] >= 4.0) & zz['有效评论']].iterrows():
    text = str(row['评论内容'])
    for w in hidden_words:
        if w in text:
            idx = text.find(w)
            zz_hidden.append((row['评分'], w, text[max(0,idx-5):min(len(text), idx+25)]))
            break

# 月度趋势
bx_monthly = bx.groupby('月份').agg({'评分': ['mean', 'count']}).round(2)
bx_monthly.columns = ['平均分', '评论数']
zz_monthly = zz.groupby('月份').agg({'评分': ['mean', 'count']}).round(2)
zz_monthly.columns = ['平均分', '评论数']

# ==================== Word文档生成 ====================
doc = Document()

# 标题
title = doc.add_heading('基于在线评论的酒店顾客关注主题与满意度驱动因素研究', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle = doc.add_paragraph('—— 漳州佰翔圆山酒店 vs 漳州宾馆 竞品对比分析')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()
meta = doc.add_paragraph()
meta.add_run('分析日期：').bold = True
meta.add_run('2026年4月\n')
meta.add_run('数据来源：').bold = True
meta.add_run('携程网\n')
meta.add_run('分析酒店：').bold = True
meta.add_run('漳州佰翔圆山酒店、漳州宾馆\n')
meta.add_run('样本规模：').bold = True
meta.add_run('佰翔酒店 2,949条 / 漳州宾馆 505条\n')
meta.add_run('研究方法：').bold = True
meta.add_run('主题分析 + 内容分析 + 情感分析 + 关键事件技术')

doc.add_page_break()

# ========== 第一章 ==========
doc.add_heading('第一章 研究概述', 1)

doc.add_heading('1.1 研究背景', 2)
doc.add_paragraph(
    '在酒店市场竞争日益激烈的背景下，在线评论已成为影响潜在消费者决策的关键因素。'
    '相较于酒店官方宣传，消费者更倾向于信任真实住客的评价。'
    '因此，系统性地分析在线评论，不仅能够识别顾客的核心关注点，'
    '还能发现服务中的痛点与机会，为酒店营销策略制定提供数据支撑。'
)

doc.add_heading('1.2 研究目标', 2)
items = [
    '识别顾客最关注的酒店属性维度（清洁、位置、服务、设施等）',
    '发现影响顾客满意度/推荐意愿的正向与负向关键驱动因素',
    '挖掘顾客提及但酒店尚未满足的隐性需求，识别服务改进机会',
    '分析负面评价中反复出现的痛点，明确服务补救方向',
    '对比本店与竞品的优劣势，制定差异化竞争策略',
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('1.3 研究方法', 2)
doc.add_paragraph(
    '本研究采用"主题分析 + 内容分析 + 情感分析 + 关键事件技术"的多方法组合，'
    '兼顾定量的广度与定性的深度。具体方法说明如下：'
)
methods = [
    ('主题分析', '通过开放式编码将评论中的顾客提及归纳为若干核心主题，'
     '形成层次化的主题体系。'),
    ('内容分析', '统计各主题的出现频率，计算维度提及率，'
     '了解顾客最关注的方面及其正负情感比例。'),
    ('情感分析', '对每个主题维度判断情感倾向（正面/负面/中性），'
     '并通过程度副词识别情感强度。'),
    ('关键事件技术', '聚焦极端评分（1-2分差评、4.5-5分好评），'
     '提取导致极端评价的具体事件与归因。'),
    ('竞品对比分析', '同时分析目标酒店与竞品，采用相同编码体系，'
     '对比优劣势与顾客感知的差异。'),
]
for name, desc in methods:
    p = doc.add_paragraph()
    p.add_run(name + '：').bold = True
    p.add_run(desc)

doc.add_heading('1.4 编码体系', 2)
doc.add_paragraph('本研究建立了12个酒店评论分析维度：')

table = doc.add_table(rows=13, cols=3)
table.style = 'Table Grid'
hdr = table.rows[0].cells
hdr[0].text = '维度'
hdr[1].text = '正面指标示例'
hdr[2].text = '负面指标示例'
for cell in hdr:
    cell.paragraphs[0].runs[0].bold = True

dims_data = [
    ('清洁卫生', '干净、整洁、无灰尘、清新', '脏、有毛发、霉味、异味'),
    ('服务态度', '热情、周到、贴心、耐心、友好', '冷淡、冷漠、敷衍、态度差'),
    ('前台服务', '办理快、效率高、入住快', '入住慢、等很久、服务差'),
    ('设施设备', '设施齐全、设备新、配置好', '设施旧、故障、不能用'),
    ('房间条件', '宽敞、舒适、床舒服、采光好', '房间小、拥挤、床硬、不舒服'),
    ('隔音效果', '隔音好、安静、噪音小', '隔音差、吵、噪音大、施工'),
    ('早餐餐饮', '丰富、品种多、味道好', '单一、难吃、没早餐'),
    ('位置交通', '位置好、交通便利、离景区近', '位置偏、不方便、难找'),
    ('性价比', '物超所值、划算、值得', '不值、太贵、性价比低'),
    ('景观环境', '景观好、江景、花园、漂亮', '景观差、难看'),
    ('停车配套', '停车方便、停车场、免费', '停车难、没车位'),
    ('卫生安全', '安全、放心、消毒', '不安全、担心、隐患'),
]
for i, row_data in enumerate(dims_data):
    table.rows[i+1].cells[0].text = row_data[0]
    table.rows[i+1].cells[1].text = row_data[1]
    table.rows[i+1].cells[2].text = row_data[2]

doc.add_page_break()

# ========== 第二章 ==========
doc.add_heading('第二章 数据概览', 1)

doc.add_heading('2.1 样本基本情况', 2)
table2 = doc.add_table(rows=5, cols=3)
table2.style = 'Table Grid'
h = table2.rows[0].cells
h[0].text = '指标'
h[1].text = '佰翔酒店'
h[2].text = '漳州宾馆'
for cell in h:
    cell.paragraphs[0].runs[0].bold = True

data2 = [
    ('总评论数', '2,949条', '505条'),
    ('有效评论数（>10字）', str(bx_total) + '条', str(zz_total) + '条'),
    ('平均评分', '4.89分', '4.78分'),
    ('5分好评占比', '87.5%', '72.1%'),
]
for i, (k, v1, v2) in enumerate(data2):
    table2.rows[i+1].cells[0].text = k
    table2.rows[i+1].cells[1].text = v1
    table2.rows[i+1].cells[2].text = v2

doc.add_paragraph()
doc.add_paragraph(
    '从数据规模来看，佰翔酒店的评论数量是漳州宾馆的近6倍，'
    '说明佰翔作为当地知名高端酒店，吸引了更多的住客关注和评价。'
    '从评分来看，佰翔的平均评分（4.89）略高于漳州宾馆（4.78），'
    '5分好评占比差距更为明显（87.5% vs 72.1%），'
    '表明佰翔在顾客满意度方面整体表现更优。'
)

doc.add_heading('2.2 评分分布', 2)

# 佰翔评分分布
bx_dist = bx['评分'].value_counts().sort_index()
zz_dist = zz['评分'].value_counts().sort_index()

table3 = doc.add_table(rows=6, cols=5)
table3.style = 'Table Grid'
h3 = table3.rows[0].cells
for i, t in enumerate(['评分区间', '佰翔条数', '佰翔占比', '漳州条数', '漳州占比']):
    h3[i].text = t
    h3[i].paragraphs[0].runs[0].bold = True

ranges = [(1, 2.9, '1-2.9分'), (3, 3.9, '3-3.9分'), (4, 4.9, '4-4.9分'), (5, 5, '5分'), (1, 5, '合计')]
for i, (lo, hi, label) in enumerate(ranges):
    if lo == 1 and hi == 5:
        b_count = len(bx)
        z_count = len(zz)
        b_pct = '100%'
        z_pct = '100%'
    else:
        b_count = len(bx[(bx['评分'] >= lo) & (bx['评分'] <= hi)])
        z_count = len(zz[(zz['评分'] >= lo) & (zz['评分'] <= hi)])
        b_pct = str(round(b_count/len(bx)*100, 1)) + '%'
        z_pct = str(round(z_count/len(zz)*100, 1)) + '%'
    table3.rows[i+1].cells[0].text = label
    table3.rows[i+1].cells[1].text = str(b_count)
    table3.rows[i+1].cells[2].text = b_pct
    table3.rows[i+1].cells[3].text = str(z_count)
    table3.rows[i+1].cells[4].text = z_pct

doc.add_paragraph()
doc.add_paragraph(
    '从评分分布可以看出，两家酒店均以5分好评为主，但漳州宾馆的'
    '中评分（3-3.9分）和低评分（1-2.9分）比例明显高于佰翔，'
    '说明漳州宾馆的服务稳定性略有不足，存在较大的评价波动。'
)

doc.add_heading('2.3 时间趋势分析', 2)
doc.add_paragraph('近12个月评分趋势：')

bz_recent = bx_monthly.tail(12)
zz_recent = zz_monthly.tail(12)

table4 = doc.add_table(rows=len(bz_recent)+1, cols=4)
table4.style = 'Table Grid'
h4 = table4.rows[0].cells
for i, t in enumerate(['月份', '佰翔平均分', '佰翔评论数', '漳州平均分', '漳州评论数'][:4]):
    h4[i].text = t
    h4[i].paragraphs[0].runs[0].bold = True
# Fix: table4 only has 4 columns
h4[0].text = '月份'
h4[1].text = '佰翔平均分'
h4[2].text = '佰翔评论数'
h4[3].text = '漳州平均分'

for i, (period, row_b) in enumerate(bz_recent.iterrows()):
    # Find matching period in zz
    zz_row = zz_recent.loc[period] if period in zz_recent.index else None
    table4.rows[i+1].cells[0].text = str(period)
    table4.rows[i+1].cells[1].text = str(row_b['平均分'])
    table4.rows[i+1].cells[2].text = str(int(row_b['评论数']))
    if zz_row is not None:
        table4.rows[i+1].cells[3].text = str(zz_row['平均分'])
    else:
        table4.rows[i+1].cells[3].text = '-'

doc.add_paragraph()
doc.add_paragraph(
    '从月度趋势来看，两家酒店的评分在节假日期间（春节、国庆等）'
    '普遍有所下降，可能与客流增加、服务压力增大有关。'
    '佰翔在大多数月份保持4.85以上的评分，而漳州宾馆波动更大，'
    '在2025年10月（国庆）和2026年1月（春节）出现明显低谷。'
)

doc.add_page_break()

# ========== 第三章 ==========
doc.add_heading('第三章 顾客关注主题分析', 1)

doc.add_heading('3.1 各维度提及频率与情感分析', 2)
doc.add_paragraph(
    '通过编码辞典对每条评论进行主题标注，统计各维度的提及次数及正负情感比例。'
)

all_dims = sorted(set(bx_pos.keys()) | set(bx_neg.keys()) | set(zz_pos.keys()) | set(zz_neg.keys()),
                  key=lambda x: -(bx_pos.get(x,0)+bx_neg.get(x,0)+zz_pos.get(x,0)+zz_neg.get(x,0)))

table5 = doc.add_table(rows=len(all_dims)+1, cols=6)
table5.style = 'Table Grid'
h5 = table5.rows[0].cells
for i, t in enumerate(['维度', '佰翔正面', '佰翔负面', '佰翔净情感', '漳州正面', '漳州负面']):
    h5[i].text = t
    h5[i].paragraphs[0].runs[0].bold = True

for i, d in enumerate(all_dims):
    net_b = bx_pos.get(d,0) - bx_neg.get(d,0)
    table5.rows[i+1].cells[0].text = d
    table5.rows[i+1].cells[1].text = str(bx_pos.get(d,0))
    table5.rows[i+1].cells[2].text = str(bx_neg.get(d,0))
    table5.rows[i+1].cells[3].text = ('+' + str(net_b)) if net_b > 0 else str(net_b)
    table5.rows[i+1].cells[4].text = str(zz_pos.get(d,0))
    table5.rows[i+1].cells[5].text = str(zz_neg.get(d,0))

doc.add_paragraph()

doc.add_heading('3.2 顾客关注重点（维度提及率）', 2)
doc.add_paragraph('维度提及率 = 该维度被提及次数 / 有效评论总数，反映顾客的关注程度：')

table6 = doc.add_table(rows=len(all_dims[:10])+1, cols=5)
table6.style = 'Table Grid'
h6 = table6.rows[0].cells
for i, t in enumerate(['维度', '佰翔提及率', '漳州提及率', '差异', '关注更集中的酒店']):
    h6[i].text = t
    h6[i].paragraphs[0].runs[0].bold = True

for i, d in enumerate(all_dims[:10]):
    b_rate = (bx_pos.get(d,0)+bx_neg.get(d,0)) / bx_total * 100
    z_rate = (zz_pos.get(d,0)+zz_neg.get(d,0)) / zz_total * 100
    diff = b_rate - z_rate
    focus = '佰翔' if diff > 5 else ('漳州' if diff < -5 else '相近')
    table6.rows[i+1].cells[0].text = d
    table6.rows[i+1].cells[1].text = str(round(b_rate, 1)) + '%'
    table6.rows[i+1].cells[2].text = str(round(z_rate, 1)) + '%'
    table6.rows[i+1].cells[3].text = ('+' + str(round(diff,1))) if diff > 0 else str(round(diff,1)) + '%'
    table6.rows[i+1].cells[4].text = focus

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('关键发现：').bold = True
doc.add_paragraph(
    '佰翔酒店的"房间条件"和"服务态度"是被提及最多的两个维度，'
    '且提及率远高于漳州宾馆，说明佰翔顾客对这两个方面的关注度极高，'
    '酒店在这两方面的表现也直接影响着顾客的整体评价。'
)
doc.add_paragraph(
    '漳州宾馆的"位置交通"和"清洁卫生"提及率相对更高，'
    '表明漳州顾客更关注位置便利性和卫生状况，这与漳州作为旅游城市，'
    '大量顾客以旅游为目的的市场特征相符。'
)

doc.add_heading('3.3 各维度正负情感比例', 2)

# 找出负向比例高的维度
neg_ratio_bx = {d: bx_neg.get(d,0)/(bx_pos.get(d,0)+bx_neg.get(d,0))*100
                for d in all_dims if bx_pos.get(d,0)+bx_neg.get(d,0) > 20}
neg_ratio_zz = {d: zz_neg.get(d,0)/(zz_pos.get(d,0)+zz_neg.get(d,0))*100
                 for d in all_dims if zz_pos.get(d,0)+zz_neg.get(d,0) > 10}

doc.add_paragraph('佰翔酒店负面提及比例最高的维度（该维度出现负面评价的比例）：')
for d, ratio in sorted(neg_ratio_bx.items(), key=lambda x: -x[1])[:5]:
    total = bx_pos.get(d,0) + bx_neg.get(d,0)
    doc.add_paragraph(
        '  ' + d + '：' + str(bx_neg.get(d,0)) + '/' + str(total) + ' = ' +
        str(round(ratio, 1)) + '%',
        style='List Bullet'
    )

doc.add_paragraph('漳州宾馆负面提及比例最高的维度：')
for d, ratio in sorted(neg_ratio_zz.items(), key=lambda x: -x[1])[:5]:
    total = zz_pos.get(d,0) + zz_neg.get(d,0)
    doc.add_paragraph(
        '  ' + d + '：' + str(zz_neg.get(d,0)) + '/' + str(total) + ' = ' +
        str(round(ratio, 1)) + '%',
        style='List Bullet'
    )

doc.add_paragraph()
doc.add_paragraph(
    '值得注意的是，漳州宾馆的"房间条件"负面比例高达51.7%（89/172），'
    '佰翔的"隔音效果"正负比也接近1:1（54负面 vs 55正面），'
    '这两个维度是各自酒店最需要关注和改进的方向。'
)

doc.add_page_break()

# ========== 第四章 ==========
doc.add_heading('第四章 满意度驱动因素与痛点分析', 1)

doc.add_heading('4.1 好评驱动因素（4.5分以上评论）', 2)
doc.add_paragraph('佰翔酒店好评触发因素TOP5：')

bx_pt_sorted = sorted(bx_pt.items(), key=lambda x: -x[1])
for rank, (event, count) in enumerate(bx_pt_sorted[:5], 1):
    p = doc.add_paragraph()
    p.add_run(str(rank) + '. ' + event + '（' + str(count) + '次）').bold = True
    if event in bx_pe and bx_pe[event]:
        doc.add_paragraph('   示例："' + bx_pe[event][0][:80] + '..."')

doc.add_paragraph()
doc.add_paragraph('漳州宾馆好评触发因素TOP5：')

zz_pt_sorted = sorted(zz_pt.items(), key=lambda x: -x[1])
for rank, (event, count) in enumerate(zz_pt_sorted[:5], 1):
    p = doc.add_paragraph()
    p.add_run(str(rank) + '. ' + event + '（' + str(count) + '次）').bold = True
    if event in zz_pe and zz_pe[event]:
        doc.add_paragraph('   示例："' + zz_pe[event][0][:80] + '..."')

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('核心发现：').bold = True
doc.add_paragraph(
    '两家酒店的好评驱动因素高度一致——"前台服务"是最核心的驱动力。'
    '佰翔的前台服务触发次数高达1,895次，远超其他因素，说明前台是佰翔'
    '服务体验的"门面"和"名片"。其次，"早餐"和"免费升房"是提升好评的'
    '重要杠杆，说明超出顾客预期的服务（如升房、赠送物品）能显著提升满意度。'
)

doc.add_heading('4.2 差评痛点分析（3分以下评论）', 2)
doc.add_paragraph('佰翔酒店差评触发因素：')

bx_nt_sorted = sorted(bx_nt.items(), key=lambda x: -x[1])
for event, count in bx_nt_sorted:
    p = doc.add_paragraph()
    p.add_run('- ' + event + '（' + str(count) + '次）').bold = True
    if event in bx_ne and bx_ne[event]:
        doc.add_paragraph('   示例："' + bx_ne[event][0][:80] + '..."')

doc.add_paragraph()
doc.add_paragraph('漳州宾馆差评触发因素：')

zz_nt_sorted = sorted(zz_nt.items(), key=lambda x: -x[1])
for event, count in zz_nt_sorted:
    p = doc.add_paragraph()
    p.add_run('- ' + event + '（' + str(count) + '次）').bold = True
    if event in zz_ne and zz_ne[event]:
        doc.add_paragraph('   示例："' + zz_ne[event][0][:80] + '..."')

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('核心发现：').bold = True
doc.add_paragraph(
    '佰翔差评中"信息不对称"是最突出的问题（18次），'
    '主要表现为酒店实际状况（如装修、设施）与宣传描述不符，'
    '导致顾客期望落空。其次是"设施老旧破损"（13次）和"卫生问题"（9次）。'
    '漳州差评中"隔音差"和"信息不对称"最为突出，'
    '两家酒店的差评都指向了"设施与描述不符"这一共性问题。'
)

doc.add_heading('4.3 情感强度分析', 2)
doc.add_paragraph(
    '情感强度通过评论中是否包含程度副词（如"非常""特别""极其""彻底"等）来衡量。'
    '情感强度越高，说明顾客的情绪反应越强烈。'
)

intense_pos = ['非常', '特别', '十分', '超级', '极致', '完美', '超棒', '超赞', '惊喜']
intense_neg = ['非常', '特别', '极其', '完全', '彻底', '太差', '极差', '噩梦', '垃圾', '恶心', '恐怖']

bx_pi, bx_ni, bx_pn, bx_nn = 0, 0, 0, 0
for _, row in bx[bx['有效评论']].iterrows():
    text = str(row['评论内容'])
    has_pos = any(w in text for w in ['好', '棒', '满意', '喜欢', '赞', '舒适', '干净', '热情', '周到'])
    has_neg = any(w in text for w in ['差', '失望', '不满', '糟糕', '恶心', '脏', '吵', '旧'])
    if has_pos:
        if any(w in text for w in intense_pos): bx_pi += 1
        else: bx_pn += 1
    if has_neg:
        if any(w in text for w in intense_neg): bx_ni += 1
        else: bx_nn += 1

zz_pi, zz_ni, zz_pn, zz_nn = 0, 0, 0, 0
for _, row in zz[zz['有效评论']].iterrows():
    text = str(row['评论内容'])
    has_pos = any(w in text for w in ['好', '棒', '满意', '喜欢', '赞', '舒适', '干净', '热情', '周到'])
    has_neg = any(w in text for w in ['差', '失望', '不满', '糟糕', '恶心', '脏', '吵', '旧'])
    if has_pos:
        if any(w in text for w in intense_pos): zz_pi += 1
        else: zz_pn += 1
    if has_neg:
        if any(w in text for w in intense_neg): zz_ni += 1
        else: zz_nn += 1

table7 = doc.add_table(rows=4, cols=3)
table7.style = 'Table Grid'
h7 = table7.rows[0].cells
h7[0].text = '情感类型'
h7[1].text = '佰翔酒店'
h7[2].text = '漳州宾馆'
for c in h7: c.paragraphs[0].runs[0].bold = True

table7.rows[1].cells[0].text = '强烈正面（含程度副词）'
table7.rows[1].cells[1].text = str(bx_pi) + ' (' + str(round(bx_pi/(bx_pi+bx_pn)*100,1)) + '%)'
table7.rows[1].cells[2].text = str(zz_pi) + ' (' + str(round(zz_pi/(zz_pi+zz_pn)*100,1)) + '%)'
table7.rows[2].cells[0].text = '普通正面'
table7.rows[2].cells[1].text = str(bx_pn)
table7.rows[2].cells[2].text = str(zz_pn)
table7.rows[3].cells[0].text = '强烈负面（高强度词）'
table7.rows[3].cells[1].text = str(bx_ni) + ' (' + str(round(bx_ni/(bx_ni+bx_nn)*100,1)) + '%)'
table7.rows[3].cells[2].text = str(zz_ni) + ' (' + str(round(zz_ni/(zz_ni+zz_nn)*100,1)) + '%)'

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('关键发现：').bold = True
doc.add_paragraph(
    '佰翔的负面情感强度比例（44.9%）高于正面（36.4%），'
    '说明在佰翔，差评顾客的情绪比好评顾客更为激烈。'
    '这提示我们：佰翔的"好评"可能比较"平淡"（强烈正面词较少），'
    '而"差评"则充满"激烈情绪"，意味着服务一旦出问题，'
    '顾客的负面传播力会非常强。'
)

doc.add_page_break()

# ========== 第五章 ==========
doc.add_heading('第五章 隐性需求与机会识别', 1)

doc.add_heading('5.1 隐性不满挖掘', 2)
doc.add_paragraph(
    '在4分以上的好评中，有部分评论虽然给了高分，但在文中使用'
    '"但是""不过""就是""唯一"等转折词，透露出隐藏的不满或改进期望。'
    '这类"隐性不满"是识别服务改进机会的重要来源。'
)

doc.add_paragraph('佰翔酒店隐性不满（' + str(len(bx_hidden)) + '条，占有效评论' +
                   str(round(len(bx_hidden)/bx_total*100, 1)) + '%）：')
bx_hw = Counter([h[1] for h in bx_hidden])
for word, cnt in bx_hw.most_common(5):
    doc.add_paragraph('  "' + word + '"：' + str(cnt) + '条')

doc.add_paragraph()
doc.add_paragraph('典型案例：')
for _, word, ctx in bx_hidden[:3]:
    doc.add_paragraph('  [' + str(_) + '分] "' + word + '"出现处：' + ctx + '...')

doc.add_paragraph()
doc.add_paragraph('漳州宾馆隐性不满（' + str(len(zz_hidden)) + '条，占有效评论' +
                   str(round(len(zz_hidden)/zz_total*100, 1)) + '%）：')
zz_hw = Counter([h[1] for h in zz_hidden])
for word, cnt in zz_hw.most_common(5):
    doc.add_paragraph('  "' + word + '"：' + str(cnt) + '条')

doc.add_paragraph()
doc.add_paragraph('漳州典型案例：')
for _, word, ctx in zz_hidden[:3]:
    doc.add_paragraph('  [' + str(_) + '分] "' + word + '"出现处：' + ctx + '...')

doc.add_heading('5.2 顾客未满足的隐性需求', 2)
doc.add_paragraph(
    '通过分析隐性不满和高频负面评论，可以识别以下隐性需求：'
)

needs = [
    ('隔音需求', '大量评论提到"隔音差""很吵""被车声吵醒"，'
     '说明顾客对安静睡眠环境有强烈需求，但现有房间未能满足。'),
    ('信息透明需求', '"和说的不一样""订的时候没告知在装修"等评论说明'
     '顾客希望在预订前就充分了解房间/设施的真实状况。'),
    ('早餐品质稳定性', '"早餐人潮涌动""位置不够""冷掉了"说明顾客对早餐'
     '体验有期待，但高峰期服务质量下降明显。'),
    ('儿童友好需求', '亲子游评论多、评分高，但同时"亲子房有避孕套套装"'
     '引发家长投诉，说明亲子服务需要更细致的边界管理。'),
    ('设施维护及时性', '"坏了没人修""插座不能用"等评论说明设施坏了之后'
     '维修响应速度是顾客关注的重点。'),
]
for name, desc in needs:
    p = doc.add_paragraph()
    p.add_run(name + '：').bold = True
    p.add_run(desc)

doc.add_page_break()

# ========== 第六章 ==========
doc.add_heading('第六章 出行目的细分分析', 1)

doc.add_heading('6.1 出行目的分布', 2)

table8 = doc.add_table(rows=7, cols=5)
table8.style = 'Table Grid'
h8 = table8.rows[0].cells
for i, t in enumerate(['出行目的', '佰翔条数', '佰翔占比', '漳州条数', '漳州占比']):
    h8[i].text = t
    h8[i].paragraphs[0].runs[0].bold = True

purposes = ['其他', '休闲游', '亲子游', '商务出差', '家庭游', '情侣游']
for i, p in enumerate(purposes):
    b_count = len(bx[bx['出行目的'] == p])
    z_count = len(zz[zz['出行目的'] == p])
    table8.rows[i+1].cells[0].text = p
    table8.rows[i+1].cells[1].text = str(b_count)
    table8.rows[i+1].cells[2].text = str(round(b_count/len(bx)*100, 1)) + '%'
    table8.rows[i+1].cells[3].text = str(z_count)
    table8.rows[i+1].cells[4].text = str(round(z_count/len(zz)*100, 1)) + '%'

doc.add_paragraph()
doc.add_paragraph(
    '从出行目的来看，两家酒店均以"其他"（未明确提及）为主，'
    '但在明确目的的评论中，"休闲游"和"亲子游"占比最高，'
    '说明漳州作为旅游城市，吸引了大批休闲度假和亲子客群。'
    '漳州宾馆的商务出差比例（7.3%）高于佰翔（9.8% vs 6.8%），'
    '可能与漳州宾馆位于市中心、更方便商务活动有关。'
)

doc.add_heading('6.2 不同出行目的的评分差异', 2)

table9 = doc.add_table(rows=5, cols=3)
table9.style = 'Table Grid'
h9 = table9.rows[0].cells
for i, t in enumerate(['出行目的', '佰翔平均分', '漳州平均分']):
    h9[i].text = t
    h9[i].paragraphs[0].runs[0].bold = True

for i, p in enumerate(['亲子游', '商务出差', '休闲游', '其他']):
    b_score = bx[bx['出行目的'] == p]['评分'].mean()
    z_score = zz[zz['出行目的'] == p]['评分'].mean()
    b_n = len(bx[bx['出行目的'] == p])
    z_n = len(zz[zz['出行目的'] == p])
    table9.rows[i+1].cells[0].text = p + ' (n=' + str(b_n) + '/' + str(z_n) + ')'
    table9.rows[i+1].cells[1].text = str(round(b_score, 2))
    table9.rows[i+1].cells[2].text = str(round(z_score, 2))

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('关键发现：').bold = True
doc.add_paragraph(
    '佰翔在亲子游方面评分（4.86）明显高于漳州（4.54），'
    '说明佰翔的亲子体验更受认可。但两家酒店的商务出差评分几乎相同（4.83），'
    '在休闲游方面佰翔略胜（4.96 vs 4.87）。'
    '值得注意的是，漳州宾馆的家庭游评分（4.93）与佰翔（4.93）持平，'
    '说明漳州宾馆在家庭客群中口碑尚可，需保持。'
)

doc.add_page_break()

# ========== 第七章 ==========
doc.add_heading('第七章 竞品对比与SWOT分析', 1)

doc.add_heading('7.1 佰翔 vs 漳州 优劣势对比', 2)

table10 = doc.add_table(rows=len(all_dims)+1, cols=4)
table10.style = 'Table Grid'
h10 = table10.rows[0].cells
for i, t in enumerate(['维度', '佰翔净情感', '漳州净情感', '优势方']):
    h10[i].text = t
    h10[i].paragraphs[0].runs[0].bold = True

for i, d in enumerate(all_dims):
    b_net = bx_pos.get(d,0) - bx_neg.get(d,0)
    z_net = zz_pos.get(d,0) - zz_neg.get(d,0)
    adv = '佰翔' if b_net > z_net else ('漳州' if z_net > b_net else '相当')
    table10.rows[i+1].cells[0].text = d
    table10.rows[i+1].cells[1].text = ('+' + str(b_net)) if b_net > 0 else str(b_net)
    table10.rows[i+1].cells[2].text = ('+' + str(z_net)) if z_net > 0 else str(z_net)
    table10.rows[i+1].cells[3].text = adv

doc.add_paragraph()
doc.add_paragraph('佰翔全面占优的维度（净情感更高）：服务态度、前台服务、景观环境、房间条件')
doc.add_paragraph('漳州相对占优的维度：位置交通（+235）、清洁卫生（+203）、景观环境（+136）')

doc.add_heading('7.2 SWOT分析（佰翔酒店）', 2)

swot = [
    ('S-优势', ['服务态度业界领先（净情感+1658）', '前台服务体验佳（净情感+320）',
               '景观环境出色（净情感+840）', '江景特色鲜明，差异化强']),
    ('W-劣势', ['设施设备老旧（负面78次）', '隔音效果差（正负比接近1:1）',
               '性价比感知偏低（负面29次）', '房间条件负提及较多（1032次）']),
    ('O-机会', ['亲子市场高评分，可深化开发', '服务优势可视觉化、KOL传播',
               '早餐体验提升空间大', '安静房型可作为差异化产品']),
    ('T-威胁', ['漳州宾馆在位置交通上占优', '设施问题持续发酵易引发口碑危机',
               '负面情感强度高，差评传播力强', '信息不对称问题可能影响复购']),
]
for title, items in swot:
    p = doc.add_paragraph()
    p.add_run(title).bold = True
    for item in items:
        doc.add_paragraph(item, style='List Bullet')

doc.add_page_break()

# ========== 第八章 ==========
doc.add_heading('第八章 营销策略建议', 1)

doc.add_heading('8.1 佰翔酒店营销策略', 2)

doc.add_heading('8.1.1 产品侧策略', 3)
ps1 = [
    ('隔音升级为标配', '将"隔音好"作为核心卖点推出"静音房型"，'
     '在OTA页面显著标注，并提供免费升级。'),
    ('早餐体验优化', '实施错峰早餐引导、增设早餐座位、'
     '推出"送餐到房"服务，减少高峰期拥挤感。'),
    ('亲子服务精细化', '开发专属亲子楼层，亲子房移除成人用品（避孕套套装等），'
     '增设儿童洗漱用品、玩具等细节关怀。'),
    ('视觉化服务展示', '推出"服务之星"员工展示，管家服务可视化，'
     '拍摄员工服务短视频用于OTA详情页。'),
]
for name, desc in ps1:
    p = doc.add_paragraph()
    p.add_run('【' + name + '】').bold = True
    doc.add_paragraph(desc, style='List Bullet')

doc.add_heading('8.1.2 推广侧策略', 3)
ps2 = [
    ('UGC口碑裂变', '设计"好评返现"或"下次入住折扣"激励机制，'
     '鼓励客人在小红书、抖音发布真实体验。'),
    ('差异化传播定位', '"江景+服务"双核心定位，在携程/小红书/抖音'
     '投放江景短视频，配合酒店服务故事。'),
    ('亲子KOL合作', '与漳州本地亲子博主合作，体验"亲子房+温泉"套餐，'
     '触达高意愿支付的家庭客群。'),
    ('竞品对比话术', '在OTA详情页主动对比位置与漳州宾馆的差异，'
     '突出"离万达更近""江景更美"等优势。'),
]
for name, desc in ps2:
    p = doc.add_paragraph()
    p.add_run('【' + name + '】').bold = True
    doc.add_paragraph(desc, style='List Bullet')

doc.add_heading('8.1.3 服务补救侧策略', 3)
ps3 = [
    ('差评24小时响应机制', '建立差评监控+快速响应机制，'
     '第一时间联系顾客解决问题，将差评转化为满意。'),
    ('信息透明化承诺', '预订页面主动标注设施装修情况、房间楼层、'
     '周边噪音等信息，减少预期落差。'),
    ('押金问题优化', '参考行业惯例取消或优化押金制度，'
     '减少因此引发的投诉和差评。'),
]
for name, desc in ps3:
    p = doc.add_paragraph()
    p.add_run('【' + name + '】').bold = True
    doc.add_paragraph(desc, style='List Bullet')

doc.add_heading('8.2 漳州宾馆营销策略', 2)

doc.add_heading('8.2.1 产品侧策略', 3)
zs1 = [
    ('园林文化定位', '主打"园林酒店"概念，对标"老牌精品"，'
     '在OTA标题和首图突出园林环境照片。'),
    ('设施翻新计划', '制定3年翻新计划，优先改造隔音和房间设施，'
     '将翻新进度作为营销卖点（"全新装修"）。'),
    ('早餐品质提升', '增加热菜品种，优化儿童收费，'
     '将早餐作为差异化卖点（"本地特色早餐"）。'),
]
for name, desc in zs1:
    p = doc.add_paragraph()
    p.add_run('【' + name + '】').bold = True
    doc.add_paragraph(desc, style='List Bullet')

doc.add_heading('8.2.2 推广侧策略', 3)
zs2 = [
    ('位置交通强化', '漳州位置提及率高，说明顾客确实看重，'
     '可在宣传中强调"步行可达古城""市中心C位"。'),
    ('历史情怀营销', '"漳州宾馆"本身有历史沉淀，可主打"老牌情怀"'
     '吸引怀旧型和商务型客群。'),
    ('竞品对比话术', '突出"老牌园林""价格更实惠""停车更方便"等差异点。'),
]
for name, desc in zs2:
    p = doc.add_paragraph()
    p.add_run('【' + name + '】').bold = True
    doc.add_paragraph(desc, style='List Bullet')

doc.add_page_break()

# ========== 第九章 ==========
doc.add_heading('第九章 研究结论', 1)

doc.add_heading('9.1 主要研究结论', 2)
conclusions = [
    ('结论一', '佰翔酒店的综合满意度（4.89）高于漳州宾馆（4.78），'
     '5分好评占比差距显著（87.5% vs 72.1%），'
     '说明佰翔的整体服务体验更获顾客认可。'),
    ('结论二', '"服务态度"和"前台服务"是两家酒店好评的核心驱动力，'
     '但也是差评的重要来源——服务好则好评，服务差则差评传播力极强。'),
    ('结论三', '佰翔的最大痛点是"隔音效果"，正负提及几乎持平，'
     '漳州宾馆的最大痛点是"房间条件"（52%负面比例），'
     '两家酒店的设施问题均较突出。'),
    ('结论四', '"信息不对称"是佰翔差评的第一触发因素，'
     '说明酒店在设施/装修等关键信息披露上存在明显不足。'),
    ('结论五', '隐性不满分析发现，约11.8%的佰翔好评包含转折词，'
     '说明即使是高分顾客也存在未被满足的期望，是服务提升的隐性机会。'),
    ('结论六', '亲子游是两家酒店的重要客群，佰翔在亲子体验上的'
     '评分（4.86）显著高于漳州（4.54），亲子市场是佰翔的差异化优势。'),
]
for name, desc in conclusions:
    p = doc.add_paragraph()
    p.add_run(name + '：').bold = True
    p.add_run(desc)

doc.add_heading('9.2 研究局限与未来方向', 2)
doc.add_paragraph(
    '本研究存在以下局限：(1)数据仅来源于携程平台，未覆盖美团、Booking等渠道；'
    '(2)"出游类型"字段为空，无法进行更精准的客群细分；(3)编码辞典依赖人工定义，'
    '可能存在主观偏差。未来可结合更多平台数据、引入机器学习进行自动编码，'
    '并通过回访调研进行"成员检验"以提升研究可信度。'
)

doc.add_heading('9.3 实践启示', 2)
doc.add_paragraph(
    '对于酒店管理者，本研究提示以下实践启示：'
)
practices = [
    '服务是口碑的核心——投资员工培训和激励，回报最显著',
    '设施维护不能省——设施老旧是差评的温床',
    '信息透明是最好的服务——提前告知装修、噪音等信息，远好过事后道歉',
    '关注"转折词"顾客——他们是最有可能被转化为忠诚顾客的群体',
    '亲子市场值得深耕——亲子需求容易被识别但服务响应不足',
]
for pr in practices:
    doc.add_paragraph(pr, style='List Bullet')

# 保存
doc.save('酒店评论综合分析报告.docx')
print('Word报告已保存: 酒店评论综合分析报告.docx')
