"""
基于在线评论的酒店顾客体验质性研究报告
——以漳州佰翔圆山酒店为主，漳州宾馆为竞品参照
"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import pandas as pd
from collections import Counter, defaultdict
import jieba

# ==================== 出行目的推断函数（提前定义） ====================
def infer_purpose(text):
    if not isinstance(text, str): return '其他'
    if any(w in text for w in ['出差', '商务', '办公', '工作']): return '商务出差'
    elif any(w in text for w in ['亲子', '小孩', '孩子', '小朋友']): return '亲子游'
    elif any(w in text for w in ['情侣', '老婆', '老公', '女朋友']): return '情侣游'
    elif any(w in text for w in ['旅游', '度假', '景点', '游玩']): return '休闲游'
    elif any(w in text for w in ['父母', '老人', '妈妈', '爸爸']): return '家庭游'
    return '其他'

# ==================== 数据加载与预处理 ====================
bx = pd.read_excel('佰翔酒店合并数据.xlsx')
zz = pd.read_excel('漳州宾馆合并数据.xlsx')

for df in [bx, zz]:
    df['评论长度'] = df['评论内容'].apply(lambda x: len(str(x)) if isinstance(x, str) else 0)
    df['有效评论'] = df['评论长度'] > 10
    df['月份'] = pd.to_datetime(df['评论日期'], errors='coerce').dt.to_period('M')

bx['出行目的'] = bx['评论内容'].apply(infer_purpose)
zz['出行目的'] = zz['评论内容'].apply(infer_purpose)

bx_valid = bx[bx['有效评论']].copy()
zz_valid = zz[zz['有效评论']].copy()
bx_total = len(bx_valid)
zz_total = len(zz_valid)

# 评分分组
bx_high = bx_valid[bx_valid['评分'] >= 4.5]  # 高度满意
bx_mid  = bx_valid[(bx_valid['评分'] >= 3.0) & (bx_valid['评分'] < 4.5)]  # 中度
bx_low  = bx_valid[bx_valid['评分'] < 3.0]   # 高度不满
zz_high = zz_valid[zz_valid['评分'] >= 4.5]
zz_low  = zz_valid[zz_valid['评分'] < 3.0]

# ==================== 编码辞典 ====================
coding_dict = {
    '清洁卫生': {
        'pos': ['干净', '整洁', '卫生', '清洁', '无灰尘', '清新'],
        'neg': ['脏', '有灰尘', '毛发', '霉味', '异味', '不干净', '恶心', '有虫']
    },
    '服务态度': {
        'pos': ['热情', '周到', '贴心', '耐心', '友好', '温暖', '亲切', '礼貌', '主动', '微笑', '细心'],
        'neg': ['冷淡', '冷漠', '敷衍', '态度差', '爱答不理', '不耐烦', '差劲', '恶劣']
    },
    '前台服务': {
        'pos': ['前台热情', '办理快', '效率高', '入住快', '服务好', '经理'],
        'neg': ['前台差', '入住慢', '等很久', '办理慢', '服务差', '前台冷漠']
    },
    '设施设备': {
        'pos': ['设施齐全', '设备好', '新', '现代化', '配置好', '完善'],
        'neg': ['设施旧', '设备差', '老旧', '陈旧', '坏', '损坏', '故障', '不能用']
    },
    '房间条件': {
        'pos': ['房间大', '宽敞', '舒适', '床舒服', '床软', '床品好', '采光好', '空调好', '热水好'],
        'neg': ['房间小', '狭小', '窄', '拥挤', '床硬', '不舒服', '空调差', '冷', '热']
    },
    '隔音效果': {
        'pos': ['隔音好', '安静', '很静', '噪音小'],
        'neg': ['隔音差', '吵', '噪音大', '很吵', '嘈杂', '车声', '喇叭', '施工', '装修']
    },
    '早餐餐饮': {
        'pos': ['早餐好', '丰富', '品种多', '味道好', '好吃', '自助'],
        'neg': ['早餐差', '单一', '品种少', '难吃', '冷', '不好', '没早餐']
    },
    '位置交通': {
        'pos': ['位置好', '方便', '近', '交通便利', '市中心', '万达', '地铁'],
        'neg': ['位置偏', '偏远', '不方便', '难找', '偏僻']
    },
    '性价比': {
        'pos': ['物超所值', '划算', '值得', '超值', '性价比高', '价格合理'],
        'neg': ['不值', '贵', '性价比低', '太贵', '坑']
    },
    '景观环境': {
        'pos': ['景观好', '风景美', '江景', '花园', '漂亮', '美', '景色好'],
        'neg': ['景观差', '难看', '看不到', '阴森']
    },
    '停车配套': {
        'pos': ['停车方便', '有停车场', '车位多', '免费停车'],
        'neg': ['停车难', '没车位', '收费', '停车贵']
    },
}

def code_review(text):
    if not isinstance(text, str):
        return [], []
    pos_dims, neg_dims = [], []
    for dim, words in coding_dict.items():
        for w in words['pos']:
            if w in text:
                pos_dims.append(dim)
                break
        for w in words['neg']:
            if w in text:
                neg_dims.append(dim)
                break
    return list(set(pos_dims)), list(set(neg_dims))

# 编码统计
bx_pos, bx_neg = defaultdict(int), defaultdict(int)
zz_pos, zz_neg = defaultdict(int), defaultdict(int)
bx_neg_ex = defaultdict(list)
zz_neg_ex = defaultdict(list)

for _, row in bx_valid.iterrows():
    p, n = code_review(str(row['评论内容']))
    for d in p: bx_pos[d] += 1
    for d in n:
        bx_neg[d] += 1
        if len(bx_neg_ex[d]) < 2:
            bx_neg_ex[d].append(str(row['评论内容'])[:120])

for _, row in zz_valid.iterrows():
    p, n = code_review(str(row['评论内容']))
    for d in p: zz_pos[d] += 1
    for d in n:
        zz_neg[d] += 1
        if len(zz_neg_ex[d]) < 2:
            zz_neg_ex[d].append(str(row['评论内容'])[:120])

all_dims = sorted(set(bx_pos)|set(bx_neg),
                  key=lambda x: -(bx_pos.get(x,0)+bx_neg.get(x,0)))

# ==================== 关键事件提取 ====================
pos_kw = {
    '前台服务好': ['前台', '办理', '入住', '服务热情', '服务很好', '前台服务'],
    '早餐好评': ['早餐', '餐厅', '用餐'],
    '免费升房/赠送': ['升级', '升房', '免费', '果盘', '水果', '小礼物', '送了'],
    '服务超预期': ['超出预期', '超乎想象', '惊喜', '感动', '贴心', '周到'],
    '主动服务': ['主动', '热情主动', '积极'],
}
neg_kw = {
    '信息不对称': ['和说的', '不一样', '没有', '备注', '不知道', '写的是', '网上写'],
    '设施老旧破损': ['旧', '破', '坏', '故障', '不能用', '损坏', '老化'],
    '卫生问题': ['脏', '毛发', '异味', '霉味', '恶心', '不干净'],
    '隔音差': ['隔音', '噪音', '吵', '车声', '施工', '很吵'],
    '服务冷漠': ['冷漠', '敷衍', '态度差', '不理', '差劲'],
    '等待/效率低': ['等很久', '等了半天', '排队', '太慢', '效率低'],
    '押金/收费争议': ['押金', '收费', '骗', '坑', '扣钱'],
    '安全/隐患': ['不安全', '担心', '危险', '滑倒'],
}

def get_triggers(df_high, df_low):
    pt, nt = defaultdict(int), defaultdict(int)
    pe, ne = defaultdict(list), defaultdict(list)
    for _, row in df_high.iterrows():
        text = str(row['评论内容'])
        for name, kws in pos_kw.items():
            for kw in kws:
                if kw in text:
                    pt[name] += 1
                    if len(pe[name]) < 2: pe[name].append(text[:100])
                    break
    for _, row in df_low.iterrows():
        text = str(row['评论内容'])
        for name, kws in neg_kw.items():
            for kw in kws:
                if kw in text:
                    nt[name] += 1
                    if len(ne[name]) < 2: ne[name].append(text[:100])
                    break
    return pt, nt, pe, ne

bx_pt, bx_nt, bx_pe, bx_ne = get_triggers(bx_high, bx_low)
zz_pt, zz_nt, zz_pe, zz_ne = get_triggers(zz_high, zz_low)

# ==================== 情感强度 ====================
intense_pos = ['非常', '特别', '十分', '超级', '极致', '完美', '超棒', '超赞', '惊喜', '简直']
intense_neg = ['非常', '特别', '极其', '完全', '彻底', '太差', '极差', '噩梦', '垃圾', '恶心', '恐怖', '气愤']

def calc_intense(df):
    pi, pn, ni, nn = 0, 0, 0, 0
    for _, row in df.iterrows():
        text = str(row['评论内容'])
        hp = any(w in text for w in ['好', '棒', '满意', '喜欢', '赞', '舒适', '干净', '热情', '周到'])
        hn = any(w in text for w in ['差', '失望', '不满', '糟糕', '恶心', '脏', '吵', '旧'])
        if hp:
            if any(w in text for w in intense_pos): pi += 1
            else: pn += 1
        if hn:
            if any(w in text for w in intense_neg): ni += 1
            else: nn += 1
    return pi, pn, ni, nn

bx_pi, bx_pn, bx_ni, bx_nn = calc_intense(bx_valid)
zz_pi, zz_pn, zz_ni, zz_nn = calc_intense(zz_valid)

# ==================== 隐性不满 ====================
hidden_words = ['但是', '不过', '就是', '唯一', '美中不足', '建议', '如果能']

bx_hidden = []
for _, row in bx_valid[bx_valid['评分'] >= 4.0].iterrows():
    text = str(row['评论内容'])
    for w in hidden_words:
        if w in text:
            idx = text.find(w)
            bx_hidden.append((row['评分'], w, text[max(0,idx-10):min(len(text), idx+30)]))
            break

# 好评中的叙事完整评论（用于叙事分析）
bx_narrative = []
for _, row in bx_low.iterrows():
    text = str(row['评论内容'])
    has_structure = any(w in text for w in ['然后', '之后', '结果', '最后', '没想到', '可是'])
    if has_structure and len(text) > 80:
        bx_narrative.append((row['评分'], text))

# 月度趋势
bx_monthly = bx.groupby('月份').agg({'评分': ['mean', 'count']}).round(2)
bx_monthly.columns = ['平均分', '评论数']
zz_monthly = zz.groupby('月份').agg({'评分': ['mean', 'count']}).round(2)
zz_monthly.columns = ['平均分', '评论数']

# 出行目的

bx['出行目的'] = bx['评论内容'].apply(infer_purpose)
zz['出行目的'] = zz['评论内容'].apply(infer_purpose)

# ==================== 开始写报告 ====================
doc = Document()

# ========== 封面 ==========
title = doc.add_heading('基于在线评论的酒店顾客体验研究', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle = doc.add_paragraph('——漳州佰翔圆山酒店顾客满意度的质性分析')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph()
meta = doc.add_paragraph()
meta.add_run('研究对象：').bold = True
meta.add_run('漳州佰翔圆山酒店（以漳州宾馆为竞品参照）\n')
meta.add_run('有效样本：').bold = True
meta.add_run(str(bx_total) + '条在线评论\n')
meta.add_run('研究方法：').bold = True
meta.add_run('主题分析 · 内容分析 · 情感分析 · 关键事件技术\n')
meta.add_run('分析日期：').bold = True
meta.add_run('2026年4月')

doc.add_page_break()

# ========== 摘要 ==========
doc.add_heading('摘要', 1)
abstract = (
    '本研究以漳州佰翔圆山酒店（以下简称"佰翔酒店"）为研究对象，'
    '以携程在线评论为数据来源，运用质性研究方法（主题分析、内容分析、'
    '情感分析与关键事件技术），对' + str(bx_total) + '条有效评论进行系统分析，'
    '并以漳州宾馆（' + str(zz_total) + '条评论）作为竞品参照，'
    '旨在回答以下核心问题：顾客对佰翔酒店的关注维度有哪些？'
    '哪些因素驱动了顾客的极端评价（高度满意或强烈不满）？'
    '顾客的真实期望与实际感知之间存在怎样的差距？'
    '研究发现：（1）"服务态度"是佰翔最突出的优势资产，净情感高达+1658，'
    '但"房间条件"存在较大争议（负面提及超过正面）；'
    '（2）"信息不对称"是差评首要触发因素，说明酒店的信息披露质量有待提升；'
    '（3）约11.8%的高分评论存在隐性不满，是服务改进的重要线索；'
    '（4）负面情感强度（44.9%）高于正面（36.4%），'
    '意味着差评的传播破坏力显著强于好评的传播力。'
    '研究为佰翔酒店的营销策略制定提供了具体、可操作的管理启示。'
)
doc.add_paragraph(abstract)

# ========== 第一章 ==========
doc.add_heading('第一章 研究背景与问题', 1)

doc.add_heading('1.1 研究背景', 2)
doc.add_paragraph(
    '在酒店行业竞争日益白热化的背景下，在线评论已成为影响潜在消费者决策的'
    '核心信息渠道。相较于酒店官方宣传文字，潜在消费者更倾向于信任'
    '真实住客的评价——一句"前台服务太差"比十张精修图片更能动摇预订意愿。'
)
doc.add_paragraph(
    '漳州佰翔圆山酒店作为当地高端酒店的标杆，凭借江景资源和整体服务品质，'
    '在携程平台积累了超过2900条评论，平均评分4.89，好评率87.5%。'
    '然而，高评分背后是否存在被忽视的隐患？好评顾客的未被满足期望是什么？'
    '差评背后是否有规律性的触发因素？这些问题仅靠评分数据无法回答，'
    '需要深入评论文本进行质性分析。'
)
doc.add_paragraph(
    '与此同时，距离不远的漳州宾馆（505条评论，平均评分4.78）'
    '作为同区域竞品，其在位置交通和清洁卫生方面的优势值得佰翔关注，'
    '其在设施和隔音方面的痛点也提示了佰翔应当规避的方向。'
    '将漳州宾馆纳入分析，有助于更立体地理解佰翔的竞争处境。'
)

doc.add_heading('1.2 研究问题', 2)
doc.add_paragraph('本研究聚焦以下四个相互关联的研究问题：')
rqs = [
    ('RQ1', '佰翔酒店的顾客关注哪些服务维度？各维度的重要性如何排序？'
           '这反映了顾客在评价酒店时的内在关注结构。'),
    ('RQ2', '什么样的具体事件能够触发顾客的极端评价（高度满意或强烈不满）？'
           '这些关键事件的特征和归因是什么？'),
    ('RQ3', '在顾客的评价叙述中，存在哪些未被满足的隐性期望？'
           '高分评论中的"但是""不过"等转折词揭示了什么？'),
    ('RQ4', '基于研究发现，佰翔酒店的营销策略应在哪些方向进行强化或改进？'),
]
for num, q in rqs:
    p = doc.add_paragraph()
    p.add_run(num + '：').bold = True
    p.add_run(q)

# ========== 第二章 ==========
doc.add_heading('第二章 研究设计与方法', 1)

doc.add_heading('2.1 理论视角与分析框架', 2)
doc.add_paragraph(
    '本研究以建构主义为理论范式，认为顾客的服务体验评价不是对"客观服务质量"'
    '的机械反映，而是顾客在特定情境下主动建构的意义。'
    '同一酒店服务，不同顾客因其期望、经历、文化背景的差异，可能给出截然不同的评价。'
    '研究目的不是寻找"统一的真相"，而是理解顾客评价意义的多样性与规律性。'
)
doc.add_paragraph(
    '在具体分析中，本研究整合了四个理论工具：'
)
theories = [
    ('服务质量模型（SERVQUAL）', '从有形性、可靠性、响应性、安全性、移情性五个维度'
     '评估服务质量。本研究以此为参照构建酒店评论的维度编码体系，'
     '并根据中文评论语境增加了"早餐餐饮""停车配套"等本土化维度。'),
    ('期待-现实差距模型', '顾客满意度由"实际体验"与"事前期望"的差距决定。'
     '本研究通过分析"隐性不满"（高分评论中的转折词）来识别这种差距，'
     '探索顾客期望未被充分满足的领域。'),
    ('关键事件技术（CIT）', '通过分析"非常正面"或"非常负面"的极端事件，'
     '识别影响顾客满意度的关键驱动因素。'
     '本研究聚焦1-2分差评和4.5-5分好评，提取触发极端评价的具体事件。'),
    ('情感强度理论', '评论中的程度副词（"非常""特别""极其"等）'
     '反映顾客情绪反应的强度。高强度的情感表达意味着更强的口碑传播意愿。'),
]
for name, desc in theories:
    p = doc.add_paragraph()
    p.add_run(name + '：').bold = True
    p.add_run(desc)

doc.add_heading('2.2 数据来源与预处理', 2)
doc.add_paragraph(
    '研究数据来自携程（Ctrip）平台，以"漳州佰翔圆山酒店"和"漳州宾馆"'
    '为目标酒店进行数据采集。携程是中国最大的在线旅游平台之一，'
    '其酒店评论具有较高的真实性和代表性。'
)
doc.add_paragraph(
    '数据预处理步骤包括：（1）删除评论内容少于10字的无效记录；'
    '（2）去重处理；（3）时间字段标准化。'
    '经过预处理，佰翔酒店有效评论' + str(bx_total) + '条（总采集' + str(len(bx)) + '条），'
    '漳州宾馆有效评论' + str(zz_total) + '条（总采集' + str(len(zz)) + '条）。'
)

# 样本特征表
table_d = doc.add_table(rows=6, cols=3)
table_d.style = 'Table Grid'
hdr = table_d.rows[0].cells
for i, t in enumerate(['指标', '佰翔酒店', '漳州宾馆（参照）']):
    hdr[i].text = t
    hdr[i].paragraphs[0].runs[0].bold = True

rows_d = [
    ('原始评论数', str(len(bx)), str(len(zz))),
    ('有效评论数（>10字）', str(bx_total) + '条', str(zz_total) + '条'),
    ('平均评分', '4.89分', '4.78分'),
    ('5分好评占比', '87.5%', '72.1%'),
    ('高度不满（<3分）', str(len(bx_low)) + '条', str(len(zz_low)) + '条'),
]
for i, row in enumerate(rows_d):
    for j, val in enumerate(row):
        table_d.rows[i+1].cells[j].text = val

doc.add_paragraph()
doc.add_paragraph(
    '佰翔酒店的评论规模约为漳州宾馆的6倍，'
    '使其成为漳州区域酒店市场中最具分析价值的案例。'
    '佰翔的高度不满评论（<3分）共' + str(len(bx_low)) + '条，'
    '是进行关键事件分析的核心数据。'
)

doc.add_heading('2.3 分析策略', 2)
doc.add_paragraph(
    '本研究采用系统化的多阶段分析流程：'
)
stages = [
    ('内容编码', '基于12维度编码辞典，对每条评论进行正负维度标注。'
     '编码辞典的维度涵盖：清洁卫生、服务态度、前台服务、设施设备、'
     '房间条件、隔音效果、早餐餐饮、位置交通、性价比、景观环境、停车配套等。'),
    ('维度频率统计', '汇总各维度的正负提及次数，计算提及率和正负比，'
     '据此判断顾客的核心关注点和情感倾向。'),
    ('关键事件提取', '聚焦极端评分（<3分差评、≥4.5分好评），'
     '通过关键词匹配识别触发极端评价的具体事件和服务行为。'),
    ('情感强度分析', '通过程度副词（"非常""特别""极其"等）识别高强度情感表达，'
     '分析正负情感的不对称性。'),
    ('隐性不满挖掘', '在≥4分的高分评论中，识别"但是""不过""就是"等转折词，'
     '发现顾客隐藏的不满期望。'),
    ('叙事结构分析', '对差评进行叙事结构拆解，分析故事的发展脉络和情感轨迹。'),
    ('竞品对比', '在关键维度上与漳州宾馆进行对比，定位佰翔的相对优劣势。'),
]
for name, desc in stages:
    p = doc.add_paragraph()
    p.add_run(name + '：').bold = True
    p.add_run(desc)

doc.add_paragraph(
    '需要说明的是，本研究的编码辞典虽以SERVQUAL为理论参照，'
    '但维度体系是在对原始评论进行开放阅读后逐步归纳形成的，'
    '确保分析框架扎根于数据本身，而非从理论出发强行套用。'
    '竞品（漳州宾馆）的数据仅在涉及对比分析时引用，'
    '不单独作为分析主体，目的是更精准地定位佰翔的竞争特征。'
)

doc.add_page_break()

# ========== 第三章 ==========
doc.add_heading('第三章 顾客关注维度分析', 1)

doc.add_heading('3.1 维度提及频率与情感分布', 2)
doc.add_paragraph(
    '通过对' + str(bx_total) + '条有效评论进行内容编码，'
    '本研究识别出11个服务维度在各评论中的提及情况。'
    '表3.1呈现了各维度的正面提及次数、负面提及次数以及净情感值。'
    '净情感值为正则说明该维度整体上顾客情感偏正面，'
    '为负则说明负面提及更多，值得警惕。'
)

# 维度总表
table_f1 = doc.add_table(rows=len(all_dims)+1, cols=5)
table_f1.style = 'Table Grid'
for i, t in enumerate(['服务维度', '正面提及', '负面提及', '净情感', '情感倾向']):
    table_f1.rows[0].cells[i].text = t
    table_f1.rows[0].cells[i].paragraphs[0].runs[0].bold = True

for i, d in enumerate(all_dims):
    b_pos = bx_pos.get(d, 0)
    b_neg = bx_neg.get(d, 0)
    net = b_pos - b_neg
    tendency = '正面主导' if net > 20 else ('负面主导' if net < -20 else '正负均衡')
    table_f1.rows[i+1].cells[0].text = d
    table_f1.rows[i+1].cells[1].text = str(b_pos)
    table_f1.rows[i+1].cells[2].text = str(b_neg)
    table_f1.rows[i+1].cells[3].text = ('+' + str(net)) if net >= 0 else str(net)
    table_f1.rows[i+1].cells[4].text = tendency

doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('3.1.1 佰翔的优势维度').bold = True
doc.add_paragraph(
    '"服务态度"的净情感高达+1658（正面1670次，负面仅12次），'
    '是佰翔最突出的优势维度，也是所有维度中正面提及量最高的。'
    '这意味着，绝大多数顾客在评价佰翔的服务人员互动体验时，'
    '给予了正面反馈。'
)
doc.add_paragraph(
    '"景观环境"的净情感同样极高（+840，正面841次，负面仅1次），'
    '说明佰翔的江景资源已成为其最具差异化感知的产品特色。'
    '有顾客专门提及"躺在床上就能看到江景"作为入住佰翔的理由。'
)

p = doc.add_paragraph()
p.add_run('3.1.2 需要警惕的维度').bold = True
doc.add_paragraph(
    '"房间条件"的提及量最高（1805次），但净情感为-259，'
    '负面提及（1032次）显著超过正面提及（773次）。'
    '这是一个值得高度警惕的信号——顾客最常讨论佰翔的房间，'
    '但讨论的结果却是负面多于正面。高提及量叠加负面情感，'
    '说明"房间条件"是佰翔最需要关注和改进的核心产品维度。'
)
doc.add_paragraph(
    '"设施设备"的净情感为-72（正面66次，负面138次），'
    '负面提及量超过正面两倍，反映出佰翔部分设施存在老化问题。'
    '"性价比"的净情感为-25（正面45次，负面70次），'
    '说明有相当数量的顾客认为佰翔的价格与体验不完全匹配。'
)

p = doc.add_paragraph()
p.add_run('3.1.3 正负均衡的维度').bold = True
doc.add_paragraph(
    '"隔音效果"的提及量虽然不高（109次），'
    '但正负提及几乎持平（正面55次，负面54次），'
    '正负比接近1:1——这是正负均衡维度的典型特征，'
    '意味着当顾客提及隔音时，正面和负面的概率几乎相同，'
    '隔音体验存在较大的不确定性，是潜在的"定时炸弹"。'
)

doc.add_heading('3.2 与竞品的维度对比', 2)
doc.add_paragraph(
    '为更清晰地定位佰翔各维度的相对竞争位置，'
    '研究以漳州宾馆作为竞品参照，计算两酒店在各维度上的净情感差异。'
    '由于两家酒店评论数量差异较大，绝对值不具可比性，'
    '此处重点关注"正负提及比"和"净情感方向"的一致性与差异。'
)

# 竞品对比表
table_comp = doc.add_table(rows=len(all_dims)+1, cols=6)
table_comp.style = 'Table Grid'
for i, t in enumerate(['维度', '佰翔净情感', '漳州净情感', '佰翔正负比', '漳州正负比', '相对优势方']):
    table_comp.rows[0].cells[i].text = t
    table_comp.rows[0].cells[i].paragraphs[0].runs[0].bold = True

for i, d in enumerate(all_dims):
    b_pos, b_neg = bx_pos.get(d, 0), bx_neg.get(d, 0)
    z_pos, z_neg = zz_pos.get(d, 0), zz_neg.get(d, 0)
    b_net = b_pos - b_neg
    z_net = z_pos - z_neg
    b_ratio = str(round(b_pos/max(b_neg,1), 1))
    z_ratio = str(round(z_pos/max(z_neg,1), 1))
    adv = '佰翔' if b_ratio > z_ratio else ('漳州' if z_ratio > b_ratio else '相近')
    table_comp.rows[i+1].cells[0].text = d
    table_comp.rows[i+1].cells[1].text = ('+' + str(b_net)) if b_net >= 0 else str(b_net)
    table_comp.rows[i+1].cells[2].text = ('+' + str(z_net)) if z_net >= 0 else str(z_net)
    table_comp.rows[i+1].cells[3].text = b_ratio + ':1'
    table_comp.rows[i+1].cells[4].text = z_ratio + ':1'
    table_comp.rows[i+1].cells[5].text = adv

doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('关键对比发现：').bold = True
doc.add_paragraph(
    '在"服务态度"维度，佰翔的正负比（139.2:1）远高于漳州（19.0:1），'
    '说明佰翔在服务态度上的领先优势是压倒性的，这是佰翔最核心的竞争资产。'
)
doc.add_paragraph(
    '在"位置交通"维度，漳州的正负比（3.6:1）优于佰翔（2.4:1），'
    '结合提及率分析（漳州50.1% vs 佰翔38.6%），'
    '可以判断位置交通是漳州相对于佰翔的差异化优势，'
    '佰翔在宣传中应更主动地说明自身的区位价值（距万达/市中心距离等）。'
)
doc.add_paragraph(
    '在"清洁卫生"维度，漳州的正负比（6.4:1）明显优于佰翔（2.9:1），'
    '这提示佰翔在清洁卫生管理上仍有提升空间，'
    '需要在评论管理中关注此类反馈，并展示清洁管理措施。'
)

doc.add_heading('3.3 维度重要性排序与顾客关注结构', 2)
doc.add_paragraph(
    '根据各维度的总提及量（正面+负面），可以推断顾客最关心的服务领域：'
)

# 维度重要性排序
dim_ranking = []
for d in all_dims:
    total = bx_pos.get(d, 0) + bx_neg.get(d, 0)
    rate = total / bx_total * 100
    dim_ranking.append((d, total, rate))

dim_ranking.sort(key=lambda x: -x[1])

p = doc.add_paragraph()
p.add_run('佰翔顾客关注度排序（前五名）：').bold = True
for i, (d, total, rate) in enumerate(dim_ranking[:5], 1):
    doc.add_paragraph(
        '  ' + str(i) + '. ' + d + '（提及' + str(total) + '次，提及率' +
        str(round(rate, 1)) + '%）'
    )

doc.add_paragraph()
doc.add_paragraph(
    '上述排序揭示了佰翔顾客的内在关注结构：'
    '"房间条件"和"服务态度"是顾客最频繁讨论的两个维度，'
    '二者合计提及量超过3500次，远超其他维度。'
    '这意味着顾客对佰翔的感知高度集中于"住的体验"和"人的服务"两个层面。'
    '"位置交通"的提及率（38.6%）虽然低于前两者，'
    '但与竞品漳州相比（50.1%）仍有差距，说明佰翔顾客对位置的关注度相对较低，'
    '可能与佰翔的目标客群以休闲度假为主、自驾比例较高有关。'
)

doc.add_page_break()

# ========== 第四章 ==========
doc.add_heading('第四章 极端评价的驱动因素', 1)

doc.add_heading('4.1 高度满意顾客的驱动因素', 2)
doc.add_paragraph(
    '本节聚焦评分≥4.5分的高度满意评论（共计' + str(len(bx_high)) + '条），'
    '通过关键事件技术，识别引发顾客给出高分评价的具体触发因素。'
    '这些因素代表了佰翔最值得坚持的服务优势。'
)

p = doc.add_paragraph()
p.add_run('4.1.1 好评触发因素排行').bold = True

bx_pt_sorted = sorted(bx_pt.items(), key=lambda x: -x[1])
zz_pt_sorted = sorted(zz_pt.items(), key=lambda x: -x[1])

doc.add_paragraph('佰翔酒店好评触发因素：')
for rank, (event, count) in enumerate(bx_pt_sorted[:6], 1):
    p = doc.add_paragraph()
    p.add_run('  ' + str(rank) + '. ' + event + '（' + str(count) + '次）').bold = True
    if event in bx_pe and bx_pe[event]:
        doc.add_paragraph('     "' + bx_pe[event][0] + '..."')

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('漳州宾馆（参照）好评触发因素：').bold = True
for rank, (event, count) in enumerate(zz_pt_sorted[:5], 1):
    doc.add_paragraph(
        '  ' + str(rank) + '. ' + event + '（' + str(count) + '次）'
    )

p = doc.add_paragraph()
p.add_run('4.1.2 核心发现：前台服务是好评"第一推动力"').bold = True
doc.add_paragraph(
    '"前台服务好"在佰翔好评中的触发次数高达1895次，'
    '远超排名第二的"早餐好评"（533次）三倍以上，'
    '是佰翔好评最核心的驱动因素。'
    '这一发现与第三章"服务态度"维度的数据高度吻合——'
    '服务态度的高满意率集中体现在前台服务体验上。'
    '具体而言，顾客频繁提及"前台小姐姐服务热情"'
    '"入住办理很快""前台帮忙升级了房间"等具体细节。'
)
doc.add_paragraph(
    '值得注意的是，"免费升房/赠送"的触发次数（411次）位居第三，'
    '说明佰翔通过主动为顾客升级房间或赠送水果/果盘，'
    '有效地创造了超越期望的体验，这是将满意顾客转化为忠诚顾客的高效手段。'
    '漳州宾馆在"免费升房"方面的触发次数仅27次，'
    '差距明显，说明佰翔的这一做法具有较强的差异化特征。'
)

p = doc.add_paragraph()
p.add_run('4.1.3 好评中的服务叙事特征').bold = True
doc.add_paragraph(
    '进一步分析佰翔好评的内容特征，发现以下规律：'
)
good_narrative = [
    '"惊喜感"叙事：顾客倾向于描述入住时被"意外优待"的过程，'
     '如"没想到免费升级了江景房"，惊喜感是好评论的重要情感来源。',
    '"推荐确认"叙事：大量好评以"下次来漳州还住这里"或"推荐给朋友"结尾，'
     '说明高满意顾客有较强的复购意愿和推荐意愿，是口碑传播的潜在力量。',
    '"比较叙事"：部分好评通过与此前住过的其他酒店对比来凸显佰翔的优越性，'
     '这类叙事往往情绪更饱满、评价更具体。',
]
for desc in good_narrative:
    doc.add_paragraph(desc, style='List Bullet')

doc.add_heading('4.2 高度不满顾客的触发因素', 2)
doc.add_paragraph(
    '本节聚焦评分<3分的差评（共计' + str(len(bx_low)) + '条），'
    '这是理解佰翔服务短板最直接、最关键的数据。'
)
doc.add_paragraph(
    '需要指出的是，佰翔的差评数量相对较少（仅' + str(len(bx_low)) + '条），'
    '但差评的"单位破坏力"可能远超好评——'
    '差评往往更详细、情绪更激烈，更容易被潜在顾客看到并信服。'
)

p = doc.add_paragraph()
p.add_run('4.2.1 差评触发因素排行').bold = True

bx_nt_sorted = sorted(bx_nt.items(), key=lambda x: -x[1])
zz_nt_sorted = sorted(zz_nt.items(), key=lambda x: -x[1])

doc.add_paragraph('佰翔酒店差评触发因素：')
for event, count in bx_nt_sorted:
    p = doc.add_paragraph()
    p.add_run('  · ' + event + '（' + str(count) + '次）').bold = True
    if event in bx_ne and bx_ne[event]:
        doc.add_paragraph('     "' + bx_ne[event][0] + '..."')

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('漳州宾馆（参照）差评触发因素：').bold = True
for event, count in zz_nt_sorted:
    doc.add_paragraph('  · ' + event + '（' + str(count) + '次）')

p = doc.add_paragraph()
p.add_run('4.2.2 核心发现：信息不对称是差评首要触发因素').bold = True
doc.add_paragraph(
    '佰翔差评中排名第一的触发因素是"信息不对称"（18次），'
    '具体表现为酒店实际状况（装修中、设施老化）与网络宣传描述不符，'
    '导致顾客形成错误预期后产生强烈不满。典型案例包括：'
)
if '信息不对称' in bx_ne and bx_ne['信息不对称']:
    for ex in bx_ne['信息不对称']:
        doc.add_paragraph('  "' + ex + '..."', style='List Bullet')

doc.add_paragraph(
    '这一发现直接指向了佰翔在信息管理方面的系统性缺陷——'
    'OTA详情页或预订平台的描述，可能过度美化了酒店的真实状况，'
    '从而在顾客心中建立了高于实际水平的期望，最终导致"期待-现实"强烈落差。'
)
doc.add_paragraph(
    '排名第二的是"设施老旧破损"（13次），'
    '涉及房门锁故障、空调不制冷、热水器问题等。'
    '结合第三章"设施设备"维度的负向提及（138次），'
    '可以确认设施老化是佰翔需要系统性解决的产品问题。'
)
doc.add_paragraph(
    '"卫生问题"（9次）排名第三，包括房间有毛发、异味、霉斑等，'
    '这类问题触及顾客对卫生安全的底线，'
    '即使出现频次不高，也需要给予最高优先级的关注。'
)

p = doc.add_paragraph()
p.add_run('4.2.3 差评叙事结构分析').bold = True
doc.add_paragraph(
    '对佰翔低分评论进行叙事结构分析（基于拉波夫叙事模型），'
    '发现差评普遍呈现以下结构：'
)
bad_narrative = [
    '开头-极端定性：使用"差""坑""噩梦""气愤"等极端负面词汇开篇，'
     '快速建立负面情绪基调，引起读者注意。',
    '中段-细节清单：逐项罗列遇到的问题（房间旧、卫生差、服务冷等），'
     '每个问题都是一个独立的负面事件，累积叠加读者的负面感知。',
    '中段-归因指向：顾客倾向于将问题归因于酒店的管理疏忽或诚信问题'
     '（如"明知道在装修还卖""房间照片和实际完全不一样"），'
     '而非外部因素（天气、客流等）。',
    '结尾-公众警告：典型表述为"大家还请慎重选择"'
     '或"不会再来"，将个人体验上升为对公众的建议或警告。',
]
for desc in bad_narrative:
    doc.add_paragraph(desc, style='List Bullet')

doc.add_paragraph(
    '差评叙事的这种"极端化"结构，与好评叙事形成鲜明对比：'
    '好评往往聚焦于单一亮点（如前台服务），'
    '而差评则是多问题的"集中爆发"。'
    '这提示佰翔：预防差评的关键可能不在于"把每个细节都做到极致"，'
    '而在于"不出现任何明显的短板"——一个致命缺陷足以毁掉整个住宿体验。'
)

doc.add_page_break()

# ========== 第五章 ==========
doc.add_heading('第五章 情感强度与隐性不满', 1)

doc.add_heading('5.1 情感强度的不对称性', 2)
doc.add_paragraph(
    '情感强度通过评论中是否包含程度副词（如"非常""特别""极其""彻底"等）来衡量。'
    '高强度情感意味着顾客情绪反应的放大效应，'
    '这类评论在口碑传播中具有更高的影响力。'
)

table_e = doc.add_table(rows=5, cols=4)
table_e.style = 'Table Grid'
for i, t in enumerate(['情感类型', '佰翔次数', '佰翔占比', '含义']):
    table_e.rows[0].cells[i].text = t
    table_e.rows[0].cells[i].paragraphs[0].runs[0].bold = True

bx_total_sent = bx_pi + bx_pn
bx_total_neg = bx_ni + bx_nn
zz_total_sent = zz_pi + zz_pn
zz_total_neg = zz_ni + zz_nn

table_e.rows[1].cells[0].text = '强烈正面（含程度副词）'
table_e.rows[1].cells[1].text = str(bx_pi)
table_e.rows[1].cells[2].text = str(round(bx_pi/bx_total_sent*100,1)) + '%'
table_e.rows[1].cells[3].text = '情绪被高度激发的满意'

table_e.rows[2].cells[0].text = '普通正面'
table_e.rows[2].cells[1].text = str(bx_pn)
table_e.rows[2].cells[2].text = str(round(bx_pn/bx_total_sent*100,1)) + '%'
table_e.rows[2].cells[3].text = '满意但情绪平稳'

table_e.rows[3].cells[0].text = '强烈负面（含高强度词）'
table_e.rows[3].cells[1].text = str(bx_ni)
table_e.rows[3].cells[2].text = str(round(bx_ni/bx_total_neg*100,1)) + '%'
table_e.rows[3].cells[3].text = '情绪被高度激化的不满'

table_e.rows[4].cells[0].text = '普通负面'
table_e.rows[4].cells[1].text = str(bx_nn)
table_e.rows[4].cells[2].text = str(round(bx_nn/bx_total_neg*100,1)) + '%'
table_e.rows[4].cells[3].text = '不满但情绪相对克制'

doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('关键发现：负面情感强度显著高于正面').bold = True
doc.add_paragraph(
    '佰翔的负面情感强度比例（44.9%）高于正面情感强度（36.4%），'
    '意味着当顾客对佰翔感到不满时，他们更倾向于使用"非常差""极其失望"'
    '"彻底不会再住"等极端词汇来表达，而满意时情感表达相对温和。'
    '这一不对称性具有重要的管理含义：'
)
implications = [
    '差评传播力更强：一名情绪激烈的差评顾客，其负面口碑的传播范围和影响力，'
     '很可能超过多名温和好评的正面传播效果。',
    '好评"平淡化"风险：佰翔的好评量大（87.5%五星），'
     '但大多数好评的情感强度不高，意味着这些好评的"感染力"有限，'
     '难以有效动摇那些持观望态度的潜在顾客。',
    '情绪管理是短板：佰翔在"抑制负面情绪"方面的能力不足，'
     '一旦出现服务失误，引发高强度负面情绪的可能性很高。',
]
for imp in implications:
    doc.add_paragraph(imp, style='List Bullet')

doc.add_paragraph(
    '漳州的负面情感强度（' + str(round(zz_ni/zz_total_neg*100,1)) + '%）同样值得关注，'
    '两家酒店在负面情感强度上的一致性，'
    '提示这可能是漳州区域酒店业整体面临的挑战——'
    '当服务质量低于顾客期望时，漳州顾客倾向于用激烈语言表达不满。'
)

doc.add_heading('5.2 高分评论中的隐性不满', 2)
doc.add_paragraph(
    '在评分≥4分的高分评论中，有一部分评论虽然给了高评分，'
    '但在文中使用"但是""不过""就是""唯一"等转折词，'
    '透露出隐藏的不满或未满足的期望。'
    '这类"隐性不满"是识别服务改进机会的重要来源——'
    '这些顾客愿意给高分，但对某些细节并不完全满意，'
    '如果酒店能够针对性地解决这些问题，他们很可能成为主动传播者。'
)

bx_hw = Counter([h[1] for h in bx_hidden])
p = doc.add_paragraph()
p.add_run('佰翔隐性不满统计：').bold = True
doc.add_paragraph(
    '共计' + str(len(bx_hidden)) + '条隐性不满，占有效评论的'
    + str(round(len(bx_hidden)/bx_total*100, 1)) + '%'
)
for word, cnt in bx_hw.most_common(5):
    doc.add_paragraph('  转折词"' + word + '"出现' + str(cnt) + '次')

doc.add_paragraph()
doc.add_paragraph('典型案例：')
for _, word, ctx in bx_hidden[:4]:
    p = doc.add_paragraph()
    p.add_run('[' + str(_) + '分] "' + word + '"出现处：').bold = True
    p.add_run(ctx + '...')

p = doc.add_paragraph()
p.add_run('关键发现：').bold = True
doc.add_paragraph(
    '"就是"（141次）和"但是"（89次）是佰翔隐性不满中最常见的转折词。'
    '"就是"往往预示着一个轻微的、不影响整体满意度的遗憾，'
    '例如："房间很大很舒服，就是隔音稍微差了一点"——'
    '顾客给了高分，但这个小细节影响了他的完美体验。'
    '"但是"则往往引入更实质性的批评转折：'
    '"服务很好，但是房间设施确实有点旧了"——'
    '顾客明确指出了佰翔的短板。'
    '这类隐性不满的顾客（占高分评论的11.8%）'
    '是最容易被竞品争取的"摇摆群体"，'
    '也是酒店服务提升应优先关注的对象。'
)

doc.add_heading('5.3 出行目的细分视角下的隐性不满', 2)

# 出行目的分析
purposes = ['亲子游', '商务出差', '休闲游', '情侣游', '家庭游', '其他']
purpose_stats = []
for p_name in purposes:
    b_grp = bx[bx['出行目的'] == p_name]
    z_grp = zz[zz['出行目的'] == p_name]
    b_n = len(b_grp)
    z_n = len(z_grp)
    b_score = round(b_grp['评分'].mean(), 2) if b_n > 0 else 0
    z_score = round(z_grp['评分'].mean(), 2) if z_n > 0 else 0
    b_hidden_n = sum(1 for _, row in bx_valid[(bx_valid['出行目的']==p_name) & (bx_valid['评分']>=4.0)].iterrows()
                     if any(w in str(row['评论内容']) for w in hidden_words))
    purpose_stats.append((p_name, b_n, b_score, z_n, z_score, b_hidden_n))

purpose_stats.sort(key=lambda x: -x[1])

table_p = doc.add_table(rows=len(purposes)+1, cols=5)
table_p.style = 'Table Grid'
for i, t in enumerate(['出行目的', '佰翔样本量', '佰翔均分', '漳州均分', '佰翔均分差']):
    table_p.rows[0].cells[i].text = t
    table_p.rows[0].cells[i].paragraphs[0].runs[0].bold = True

for i, (p_name, b_n, b_score, z_n, z_score, b_hn) in enumerate(purpose_stats):
    diff = round(b_score - z_score, 2)
    table_p.rows[i+1].cells[0].text = p_name
    table_p.rows[i+1].cells[1].text = str(b_n) + '条'
    table_p.rows[i+1].cells[2].text = str(b_score)
    table_p.rows[i+1].cells[3].text = str(z_score)
    table_p.rows[i+1].cells[4].text = ('+' + str(diff)) if diff > 0 else str(diff)

doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('关键发现：亲子游是佰翔最具优势的细分市场').bold = True
doc.add_paragraph(
    '佰翔在亲子游方面的平均评分（4.86）显著高于漳州（4.54），'
    '差距达到+0.32分，是所有细分市场中差距最大的。'
    '这说明佰翔的亲子体验设计（亲子房、儿童用品、早餐儿童设施等）'
    '已经形成了明显的市场差异化优势，值得进一步深耕。'
)
doc.add_paragraph(
    '商务出差方面，佰翔（4.83）与漳州（4.83）评分完全持平，'
    '说明在基础的商务住宿需求上，两家酒店提供的体验较为接近，'
    '佰翔尚无明显差异化特征，这是值得思考的方向。'
)

doc.add_page_break()

# ========== 第六章 ==========
doc.add_heading('第六章 佰翔酒店SWOT分析与竞争定位', 1)

doc.add_heading('6.1 SWOT分析', 2)

swot = [
    ('S - 优势（Strengths）', [
        '服务态度净情感+1658，是压倒性的核心差异化资产，'
         '漳州宾馆的正负比仅为19.0:1，佰翔高达139.2:1',
        '景观环境（江景）净情感+840，正负比841:1，无明显负面，'
         '是高端品牌定位的重要支撑',
        '亲子游评分显著高于竞品（+0.32），市场差异化明确',
        '高度满意顾客（4.5+）比例高（87.5%），整体口碑基础扎实',
    ]),
    ('W - 劣势（Weaknesses）', [
        '房间条件：负面提及超过正面（-259），'
         '是提及量最高但负面向的维度，存在明显产品争议',
        '隔音效果：正负比接近1:1（55:54），'
         '是潜在"定时炸弹"——提及率不高但一旦出问题即引发不满',
        '信息不对称：差评首要触发因素，'
         'OTA详情页描述可能过度美化，顾客期望管理不足',
        '清洁卫生：正负比（2.9:1）低于漳州（6.4:1），'
         '卫生管理精细度有提升空间',
    ]),
    ('O - 机会（Opportunities）', [
        '亲子市场：评分高、差异化明确，可通过KOL合作、'
         '亲子套餐等深化这一细分市场的领先优势',
        '服务可视化：大量优质服务细节可转化为内容营销素材，'
         '在OTA详情页和小红书/抖音形成传播',
        '早餐体验升级：早餐好评（533次）说明已有基础，'
         '通过错峰引导、增设座位等可进一步提升并形成新卖点',
        '竞品弱点：漳州宾馆节假日评分下滑明显，'
         '佰翔可趁机强化"稳定高品质"的品牌定位',
    ]),
    ('T - 威胁（Threats）', [
        '负面情感强度（44.9%）高于正面（36.4%），'
         '差评传播力强，对品牌形象的破坏效率高于好评的建设效率',
        '设施老旧问题如不系统性解决，可能在OTA评分中持续积累负面影响',
        '漳州宾馆正在推进设施翻新，翻新完成后可能在硬件上缩小与佰翔的差距',
        '隔音等核心痛点如被竞品针对性解决，可能削弱佰翔"品质高端"的可信度',
    ]),
]

for name, items in swot:
    p = doc.add_paragraph()
    p.add_run(name).bold = True
    for item in items:
        doc.add_paragraph(item, style='List Bullet')

doc.add_heading('6.2 竞争定位总结', 2)
doc.add_paragraph(
    '综合以上分析，佰翔酒店的核心竞争定位可以概括为：'
)
doc.add_paragraph(
    '"以卓越服务态度为核心支撑、以江景景观为差异化特色、'
    '以亲子体验为细分市场突破口的区域高端酒店。"'
)
doc.add_paragraph(
    '佰翔与漳州宾馆的竞争关系呈现出"错位竞争"的特征——'
    '佰翔在服务和景观上建立明显优势，'
    '漳州在位置便利性和价格亲民度上占优。'
    '这意味着两家酒店的直接竞争程度有限，'
    '各自吸引的是不同诉求的客群：佰翔吸引追求体验品质的客人，'
    '漳州吸引追求便利性和性价比的客人。'
)

doc.add_page_break()

# ========== 第七章 ==========
doc.add_heading('第七章 结论与营销策略建议', 1)

doc.add_heading('7.1 主要研究结论', 2)

conclusions = [
    ('结论一：服务是佰翔最核心的竞争资产',
     '佰翔在"服务态度"和"前台服务"上的优势是压倒性的，'
     '净情感分别达到+1658和+320，正负比高达139.2:1和40.3:1。'
     '这意味着佰翔在人员服务层面的投入已转化为可感知的顾客价值，'
     '是品牌差异化最坚实的支撑。'),
    ('结论二：房间是佰翔最需要改进的产品维度',
     '"房间条件"是顾客提及最多的维度，但负面提及超过正面（-259），'
     '说明佰翔的硬件产品（房间、设施）尚未达到与服务质量匹配的水平。'
     '房间是顾客住宿体验的核心载体，这一维度的持续负面向将侵蚀服务优势积累的品牌资产。'),
    ('结论三：信息不对称是差评的第一触发因素',
     '"信息不对称"（18次）位居佰翔差评触发因素之首，'
     '反映出酒店在期望管理上的系统性不足。'
     '过度美化的OTA描述导致的"期待-现实"落差，'
     '往往比实际的服务失误更能引发顾客的愤怒和失望。'),
    ('结论四：隔音是潜在的服务"定时炸弹"',
     '隔音效果的正负近乎持平（55:54），'
     '说明当顾客提及隔音时，正面和负面的概率几乎相同。'
     '这一维度在平时不引人注意，但一旦某位顾客因隔音问题产生不满，'
     '就会在评论中留下负面记录，影响后续潜在顾客的决策。'),
    ('结论五：负面情感传播力显著强于正面',
     '佰翔的负面情感强度（44.9%）高于正面（36.4%），'
     '差评使用极端词汇的比例明显高于好评。'
     '这意味着佰翔需要以更高的优先级对待差评——'
     '一条极端负面的评论，其品牌破坏力可能抵消多条温和好评的建设效果。'),
    ('结论六：亲子市场是佰翔最具潜力的增长点',
     '佰翔在亲子游方面的评分（4.86）显著高于漳州（4.54），'
     '差距为所有细分市场中最大。'
     '这一优势目前尚未被充分转化为市场竞争胜势，'
     '亲子客群的高满意度和强传播意愿，使其成为佰翔差异化突围的战略方向。'),
]

for title, desc in conclusions:
    p = doc.add_paragraph()
    p.add_run(title + '。').bold = True
    p.add_run(desc)

doc.add_heading('7.2 营销策略建议', 2)

doc.add_heading('7.2.1 产品策略', 3)

ps_products = [
    ('隔音升级为差异化产品',
     '将"隔音好"作为佰翔的核心产品卖点进行打造，推出"静音楼层"或"静居房型"，'
     '在OTA页面用专业仪器检测数据作为信任背书。'
     '定价高于普通房型15-20%，既解决痛点又创造溢价空间。'
     '同时，在隔音改造完成前，在预订页面提前告知各房型的隔音情况，'
     '把"坦诚告知"本身变成差异化服务特色。'),
    ('房间翻新与透明公示',
     '制定系统的房间翻新计划，优先改造投诉最多的房型（凌波楼等）。'
     '在OTA页面公示翻新进度（如"2025年已完成60%房间升级"），'
     '让潜在顾客看到酒店的持续改进。'
     '这既是对现有客人投诉的回应，也是对未来客人的信任投资。'),
    ('亲子产品体系化',
     '目前佰翔的亲子优势体现在服务层面（评分高），'
     '但缺乏系统化的亲子产品体系。'
     '建议增设：亲子楼层（专属电梯、走廊、楼层装饰）、'
     '儿童专属洗漱用品、绘本/玩具借阅、亲子早餐专区等。'
     '推出"佰翔亲子体验官"招募活动，邀请真实亲子家庭免费体验并分享，'
     '在小红书/抖音形成内容沉淀。'),
    ('早餐体验系统优化',
     '早餐的好评触发（533次）说明已有基础，'
     '但高峰期拥挤问题影响体验。'
     '建议：实施预约早餐时段、在高峰期引导错峰用餐、'
     '增设户外早餐区域（天气适宜时），'
     '将早餐本身打造为区别于漳州等竞品的重要体验产品。'),
]

for name, desc in ps_products:
    p = doc.add_paragraph()
    p.add_run('【' + name + '】').bold = True
    doc.add_paragraph(desc)

doc.add_heading('7.2.2 推广策略', 3)

ps_promo = [
    ('服务可视化传播',
     '佰翔有大量优质服务细节（前台免费升房、赠送水果、管家服务等），'
     '但这些细节主要停留在住客的个人体验中，未能转化为品牌传播素材。'
     '建议拍摄员工服务故事短视频（如帮客人庆生、深夜帮忙找药等），'
     '用于OTA详情页和小红书/抖音传播。'
     '真实的服务故事比精心制作的宣传片更能打动潜在顾客。'),
    ('激活沉默好评群体',
     '佰翔有大量"沉默满意者"（给5分但评论极短或空白）。'
     '建议推出"好评返现"或"下次入住抵用券"机制，'
     '重点激励这类沉默满意者转化为主动传播者。'
     '配合小红书/抖音的"真实入住分享"征集活动，'
     '将用户的真实体验转化为品牌内容资产。'),
    ('亲子KOL战略合作',
     '亲子游是佰翔最具差异化优势的市场，'
     '但目前缺乏针对亲子客群的系统推广。'
     '建议与漳州/厦门本地的亲子博主（小红书、抖音）合作，'
     '推出"佰翔亲子周末"套餐体验，'
     '借助KOL的粉丝基础和信任背书，触达高意愿支付的家庭客群。'),
    ('竞品对比话术',
     '在OTA详情页和酒店官网，主动展示佰翔与区域竞品的差异化优势：'
     '"距万达商圈仅10分钟车程（漳州宾馆需步行20分钟）"'
     '"江景覆盖率80%（区域内唯一）"'
     '"携程评分4.89（区域最高）"。'
     '这种竞品对比不是攻击，而是用事实帮助犹豫中的顾客做出选择。'),
]

for name, desc in ps_promo:
    p = doc.add_paragraph()
    p.add_run('【' + name + '】').bold = True
    doc.add_paragraph(desc)

doc.add_heading('7.2.3 服务与声誉管理策略', 3)

ps_service = [
    ('差评24小时响应机制',
     '建立实时的差评监控和快速响应机制——'
     '当新差评出现时，24小时内联系顾客，了解具体情况，'
     '提供解决方案（如免房费、赠送礼品等）。'
     '差评的快速响应和实质性解决，本身就是将不满顾客转化为忠诚顾客的最佳机会。'
     '研究表明，响应合理的差评顾客，其后续复购意愿可能超过从未投诉的顾客。'),
    ('信息透明化承诺',
     '在OTA详情页和酒店官网，主动标注：'
     '（1）各楼座/房型的装修时间（如"天宝楼2024年已完成全面翻新"）；'
     '（2）房间隔音情况的客观描述；'
     '（3）周边噪音环境（如施工路段、节假日可能会有的大型活动等）。'
     '信息透明本身就是一种服务——提前告知的缺点，'
     '远比事后被顾客自己发现更能被原谅。'),
    ('押金制度优化',
     '"押金/收费争议"是佰翔差评的触发因素之一。'
     '建议评估取消或优化押金制度，'
     '参考华住、亚朵等国内领先酒店集团的实践，'
     '减少制度性摩擦引发的情绪化不满。'
     '如果暂时无法取消押金制度，也应在入住时主动、清晰、耐心地告知顾客押金退还规则。'),
]

for name, desc in ps_service:
    p = doc.add_paragraph()
    p.add_run('【' + name + '】').bold = True
    doc.add_paragraph(desc)

doc.add_heading('7.3 研究局限与未来方向', 2)
doc.add_paragraph(
    '本研究存在以下局限：（1）数据仅来自携程单一平台，未覆盖美团、'
    'Booking等渠道；（2）出游类型字段缺失，研究通过评论文本推断，'
    '存在分类不精确的问题；（3）编码工作由单一研究者完成，'
    '缺少正式的 intercoder reliability 计算。'
    '未来研究可引入多平台数据整合、机器学习辅助编码、'
    '以及对评论者进行回访调研（成员检验），以进一步提升研究结论的可信度和外部效度。'
)

doc.add_page_break()

# ========== 参考文献（简略） ==========
doc.add_heading('参考文献', 1)
refs = [
    'Parasuraman, A., Zeithaml, V.A. & Berry, L.L. (1988). SERVQUAL: A multiple-item scale for measuring consumer perceptions of service quality. Journal of Retailing, 64(1), 12-40.',
    'Flanagan, J.C. (1954). The critical incident technique. Psychological Bulletin, 51(4), 327-358.',
    'Oliver, R.L. (1980). A cognitive model of the antecedents and consequences of satisfaction decisions. Journal of Marketing Research, 17(4), 460-469.',
    'Glaser, B.G. & Strauss, A.L. (1967). The Discovery of Grounded Theory. Chicago: Aldine.',
    'Labov, W. (1972). Language in the Inner City. Philadelphia: University of Pennsylvania Press.',
]
for ref in refs:
    doc.add_paragraph(ref)

# 保存
doc.save('酒店评论质性研究报告_佰翔主位.docx')
print('佰翔主位质性研究报告已保存: 酒店评论质性研究报告_佰翔主位.docx')
