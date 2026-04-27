"""
报告生成模块 - 通用质性分析（QDA）多格式报告生成器
支持：Word(.docx)、Excel(.xlsx)、Markdown(.md)
输出学术风格的质性分析报告，适用于内容分析、主题编码、扎根理论等研究方法。
"""

import os
import pandas as pd
import numpy as np
from datetime import datetime
from pathlib import Path
from collections import Counter


# ==================== Word报告 ====================

def generate_word_report(results_dict, memo_manager=None, output_path=None):
    """
    生成完整Word格式质性分析报告（.docx）

    Args:
        results_dict: {分组名: {
            'df': DataFrame (原始数据),
            'stats': 基本统计结果,
            'dimension_stats': 维度/编码统计,
            'sentiment_summary': 情感摘要,
            'critical_events': 关键事件,
            'word_freq': 词频,
        }}
        memo_manager: MemoManager实例（可选）
        output_path: 输出路径

    Returns:
        str: 保存的文件路径
    """
    try:
        from docx import Document
        from docx.shared import Pt, Cm, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.oxml.ns import qn
    except ImportError:
        print('请安装: pip install python-docx')
        return None

    doc = Document()

    # 设置中文字体
    def set_run_font(run, font_name='微软雅黑', size=11, bold=False):
        run.font.name = font_name
        run.font.size = Pt(size)
        run.font.bold = bold
        run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)

    def add_heading(doc, text, level=1):
        p = doc.add_heading(text, level=level)
        for run in p.runs:
            set_run_font(run, size=16 if level == 1 else 13)
        return p

    def add_paragraph(doc, text, bold=False, indent=False):
        p = doc.add_paragraph()
        if indent:
            p.paragraph_format.left_indent = Cm(0.5)
        run = p.add_run(text)
        set_run_font(run, size=11, bold=bold)
        return p

    def add_table_row(table, data, bold=False, bg_color=None):
        row = table.add_row()
        for i, text in enumerate(data):
            cell = row.cells[i]
            cell.text = str(text)
            for para in cell.paragraphs:
                for run in para.runs:
                    set_run_font(run, size=10, bold=bold)
        return row

    group_names = list(results_dict.keys())
    group_count = len(group_names)

    # ========== 封面 ==========
    doc.add_paragraph()
    add_heading(doc, '质性分析研究报告', 0)

    p = doc.add_paragraph()
    run = p.add_run('基于文本数据的编码分析与解释性洞察')
    set_run_font(run, size=13)

    doc.add_paragraph()
    doc.add_paragraph()

    # 元信息
    meta_table = doc.add_table(rows=5, cols=2)
    meta_table.style = 'Table Grid'
    meta_data = [
        ('报告生成时间', datetime.now().strftime('%Y年%m月%d日 %H:%M')),
        ('分析分组数量', str(group_count)),
        ('编码方法', '主题编码·内容分析'),
        ('分析方法', '内容分析·情感分析·关键事件技术·扎根理论编码'),
        ('生成工具', 'QDA分析工具 v1.0'),
    ]
    for i, (k, v) in enumerate(meta_data):
        meta_table.rows[i].cells[0].text = k
        meta_table.rows[i].cells[1].text = v
        for cell in meta_table.rows[i].cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    set_run_font(run, size=10)

    doc.add_page_break()

    # ========== 第一章：项目概述 ==========
    add_heading(doc, '一、项目概述', 1)
    add_paragraph(doc,
        '本报告运用质性研究方法对文本数据进行系统分析。通过对原始文档进行编码标注、'
        '主题归纳与情感强度评估，揭示文本中蕴含的核心主题、情感倾向与关键事件。'
        '研究采用混合方法策略，将内容分析的量化统计与扎根理论的性质诠释相结合，'
        '为研究问题的回答提供多层次、多角度的证据支撑。'
    )

    add_heading(doc, '研究问题', 2)
    rqs = [
        'RQ1: 文本中最受关注的主题维度有哪些？各维度的正负情感分布如何？',
        'RQ2: 积极评价与消极评价的核心触发因素是什么？（关键事件）',
        'RQ3: 积极文本中是否隐藏着未明确表达的隐性不满？',
        'RQ4: 不同分组在各编码维度上的异同比较如何？',
    ]
    for rq in rqs:
        p = doc.add_paragraph(style='List Number')
        run = p.add_run(rq)
        set_run_font(run, size=11)

    doc.add_page_break()

    # ========== 第二章：数据来源与研究设计 ==========
    add_heading(doc, '二、数据来源与研究设计', 1)

    add_heading(doc, '2.1 数据来源', 2)
    for i, (group_name, data) in enumerate(results_dict.items()):
        stats = data.get('stats', {})
        df = data.get('df')
        label = f'项目{i+1}' if len(group_names) == len(results_dict) else group_name
        add_paragraph(doc, f'【{group_name}】', bold=True)
        t = doc.add_table(rows=4, cols=2)
        t.style = 'Table Grid'
        rows_data = [
            ('原始文档数', stats.get('total_raw', 'N/A')),
            ('有效记录数', stats.get('total_valid_record', 'N/A')),
            ('分析时间跨度', stats.get('time_span', '—')),
            ('平均文本长度（字）', round(stats.get('text_length_mean', 0)) if stats.get('text_length_mean') else 'N/A'),
        ]
        for row_i, (k, v) in enumerate(rows_data):
            t.rows[row_i].cells[0].text = k
            t.rows[row_i].cells[1].text = str(v)
        doc.add_paragraph()

    add_heading(doc, '2.2 分析策略', 2)
    strategies = [
        ('内容分析', '基于编码辞典，对文本进行结构化标注与频次统计'),
        ('情感强度分析', '区分高强度情感词与基础情感表达，识别情感极性'),
        ('关键事件技术（CIT）', '提取触发积极与消极评价的核心事件'),
        ('扎根理论编码', '从文本中归纳理论性范畴，建立维度间关系'),
        ('分组对比分析', '通过共现矩阵与交叉分析定位各组之间的异同'),
    ]
    for name, desc in strategies:
        p = doc.add_paragraph()
        run1 = p.add_run(f'· {name}：')
        set_run_font(run1, size=11, bold=True)
        run2 = p.add_run(desc)
        set_run_font(run2, size=11)

    doc.add_page_break()

    # ========== 第三章：编码系统与主题维度 ==========
    add_heading(doc, '三、编码系统与主题维度', 1)

    # 计算总维度数（跨所有分组去重）
    all_dims = set()
    for data in results_dict.values():
        dim_stats = data.get('dimension_stats', {})
        all_dims.update(dim_stats.keys())
    add_paragraph(doc,
        f'本研究采用主题编码方法，从文本中识别并归纳出{len(all_dims)}个核心编码维度。'
        '以下按分组展示各维度的编码统计结果。'
    )

    for group_name, data in results_dict.items():
        dim_stats = data.get('dimension_stats', {})
        if not dim_stats:
            continue

        add_heading(doc, f'3.{group_names.index(group_name)+1} {group_name}编码维度统计', 2)

        # 维度统计表
        sorted_dims = sorted(dim_stats.items(), key=lambda x: -(x[1].get('total', 0)))
        table = doc.add_table(rows=len(sorted_dims) + 1, cols=5)
        table.style = 'Table Grid'
        headers = ['编码维度', '正面提及', '负面提及', '净情感值', '正负比']
        for i, h in enumerate(headers):
            table.rows[0].cells[i].text = h

        for row_i, (dim, stats) in enumerate(sorted_dims, 1):
            row_data = [
                dim,
                str(stats.get('pos', 0)),
                str(stats.get('neg', 0)),
                str(stats.get('net', 0)),
                f"{stats.get('ratio', 0)}:1" if stats.get('ratio', 0) > 0 else '—'
            ]
            for col_i, val in enumerate(row_data):
                table.rows[row_i].cells[col_i].text = val

        doc.add_paragraph()

        # 关键发现
        top_dims = [d[0] for d in sorted_dims[:3]]
        p = doc.add_paragraph()
        run = p.add_run(f'关键发现：{group_name}中出现频率最高的前三位编码维度为：')
        set_run_font(run, size=11, bold=True)
        run2 = p.add_run('、'.join(top_dims))
        set_run_font(run2, size=11)

    doc.add_page_break()

    # ========== 第四章：关键事件分析 ==========
    add_heading(doc, '四、关键事件分析', 1)
    add_paragraph(doc,
        '关键事件技术（Critical Incident Technique）通过分析极端积极与极端消极的文本片段，'
        '识别驱动正向与负向评价的核心事件。本节对各分组的关键事件进行提取与分析。'
    )

    for group_name, data in results_dict.items():
        events = data.get('critical_events', {})
        add_heading(doc, f'4.{group_names.index(group_name)+1} {group_name}关键事件', 2)

        add_paragraph(doc, '【正向触发事件 Top5】', bold=True)
        pos_ev = events.get('positive_events', {})
        sorted_pos = sorted(pos_ev.items(), key=lambda x: -x[1])[:5]
        for event, count in sorted_pos:
            p = doc.add_paragraph()
            run = p.add_run(f'  · {event}：{count}次')
            set_run_font(run, size=11)

        add_paragraph(doc, '【负向触发事件 Top5】', bold=True)
        neg_ev = events.get('negative_events', {})
        sorted_neg = sorted(neg_ev.items(), key=lambda x: -x[1])[:5]
        for event, count in sorted_neg:
            p = doc.add_paragraph()
            run = p.add_run(f'  · {event}：{count}次')
            set_run_font(run, size=11)

        # 典型引文
        neg_ex = events.get('negative_examples', {})
        if neg_ex:
            add_paragraph(doc, '【负向典型引文】', bold=True)
            for event, examples in list(neg_ex.items())[:2]:
                for ex in examples[:1]:
                    p = doc.add_paragraph()
                    run = p.add_run(f'  "{ex[:120]}..."')
                    set_run_font(run, size=9)

    doc.add_page_break()

    # ========== 第五章：情感强度分析 ==========
    add_heading(doc, '五、情感强度与隐性表达分析', 1)
    add_paragraph(doc,
        '本节分析文本中的情感强度分布，识别高强度情感表达（含程度副词修饰的情感词）'
        '以及隐性负面表达（表面积极但隐含批评的文本），以揭示表面之下的深层态度。'
    )

    for group_name, data in results_dict.items():
        sent = data.get('sentiment_summary', {})
        add_paragraph(doc, f'【{group_name}】', bold=True)
        p = doc.add_paragraph()
        run = p.add_run(
            f"强正面率：{sent.get('intense_positive_rate', 0):.1f}% | "
            f"强负面率：{sent.get('intense_negative_rate', 0):.1f}% | "
            f"隐性负面率：{sent.get('hidden_dissatisfaction_rate', 0):.1f}%"
        )
        set_run_font(run, size=11)

        insight = doc.add_paragraph()
        run = insight.add_run('分析：')
        set_run_font(run, size=11, bold=True)
        hn_rate = sent.get('hidden_dissatisfaction_rate', 0)
        if hn_rate > 15:
            text = f'该分组隐性负面表达率较高（>{hn_rate:.1f}%），说明部分文本在整体积极评价的表象下存在未明确表达的负面态度，值得重点关注。'
        elif hn_rate > 8:
            text = f'该分组隐性负面表达率中等（{hn_rate:.1f}%），多数负面态度能通过直接表达体现，但仍有改进空间。'
        else:
            text = f'该分组隐性负面表达率较低（{hn_rate:.1f}%），情感表达较为直接，分析结果可信度高。'
        run2 = insight.add_run(text)
        set_run_font(run2, size=11)

    doc.add_page_break()

    # ========== 第六章：分组对比分析 ==========
    if group_count >= 2:
        add_heading(doc, '六、分组对比分析', 1)
        add_paragraph(doc,
            '通过分组间的横向比较，识别各分组在编码维度上的差异与共性，'
            '揭示不同分组在关注主题、情感倾向及关键事件上的异同。'
        )

        name_list = group_names

        d_list = [results_dict[name].get('dimension_stats', {}) for name in name_list]

        # 合并所有维度
        all_dims = set()
        for d in d_list:
            all_dims.update(d.keys())

        # 优势/差异行
        add_heading(doc, '6.1 分组间维度差异', 2)
        diff_found = False
        for dim in sorted(all_dims):
            vals = [(idx, d.get(dim, {'net': 0, 'total': 0})) for idx, d in enumerate(d_list)]
            val_strs = [f'{name_list[idx]}: 净情感{v["net"]}' for idx, v in vals]
            max_net = max(v['net'] for _, v in vals)
            min_net = min(v['net'] for _, v in vals)
            if max_net - min_net > 5:
                diff_found = True
                p = doc.add_paragraph()
                run = p.add_run(f'  · {dim} — ' + ' | '.join(val_strs))
                set_run_font(run, size=11)
        if not diff_found:
            add_paragraph(doc, '  （各分组在编码维度上差异不显著）')

        add_heading(doc, '6.2 关键事件对比', 2)
        for idx, name in enumerate(name_list):
            ev = results_dict[name].get('critical_events', {})
            pos_ev = ev.get('positive_events', {})
            neg_ev = ev.get('negative_events', {})
            top_pos = sorted(pos_ev.items(), key=lambda x: -x[1])[:3]
            top_neg = sorted(neg_ev.items(), key=lambda x: -x[1])[:3]
            p = doc.add_paragraph()
            run = p.add_run(f'{name} — 主要正向事件：')
            set_run_font(run, size=11, bold=True)
            if top_pos:
                run2 = p.add_run(', '.join(f'{e}({c})' for e, c in top_pos))
                set_run_font(run2, size=11)
            p2 = doc.add_paragraph()
            run3 = p2.add_run(f'{" "*len(name)} 主要负向事件：')
            set_run_font(run3, size=11, bold=True)
            if top_neg:
                run4 = p2.add_run(', '.join(f'{e}({c})' for e, c in top_neg))
                set_run_font(run4, size=11)

    doc.add_page_break()

    # ========== 第七章：结论 ==========
    add_heading(doc, '七、研究结论', 1)

    add_heading(doc, '7.1 主要发现', 2)
    # 基于数据自动生成结论摘要
    conclusion_items = []
    for group_name, data in results_dict.items():
        dim_stats = data.get('dimension_stats', {})
        if dim_stats:
            sorted_dims = sorted(dim_stats.items(), key=lambda x: -(x[1].get('total', 0)))
            top_dim = sorted_dims[0][0] if sorted_dims else '—'
            total_mentions = sum(v.get('total', 0) for v in dim_stats.values())
            conclusion_items.append(
                f'{group_name}的文本中，"{top_dim}"是最受关注的编码维度，'
                f'编码提及总数为{total_mentions}次。'
            )

    for i, c in enumerate(conclusion_items, 1):
        p = doc.add_paragraph(style='List Number')
        run = p.add_run(c)
        set_run_font(run, size=11)

    # 综合结论
    add_paragraph(doc,
        '综合以上分析，本研究报告揭示了文本数据中蕴含的核心主题与情感模式。'
        '通过系统性的编码分析与关键事件提取，为理解研究对象的深层特征提供了'
        '数据驱动的质性证据。研究结果可为后续的理论建构与实证研究提供参考基础。'
    )

    add_heading(doc, '7.2 研究局限与展望', 2)
    limitations = [
        '分析结果受限于文本样本的数量与质量，样本代表性有待进一步验证。',
        '编码辞典的完备性可能影响维度识别的全面性，建议结合领域知识持续优化。',
        '情感分析对反讽、隐喻等复杂修辞手法的识别仍存在局限。',
        '未来研究可结合访谈、问卷等多源数据进行三角验证。',
    ]
    for lim in limitations:
        p = doc.add_paragraph()
        run = p.add_run(f'  · {lim}')
        set_run_font(run, size=11)

    # ========== 备忘录摘要（如果有）==========
    if memo_manager:
        doc.add_page_break()
        add_heading(doc, '附录：研究者备忘录', 1)
        memos = memo_manager.export_to_dataframe()
        if not memos.empty:
            summary = memo_manager.summary()
            total_memos = summary.get('总备忘录数', 0) if isinstance(summary, dict) else 0
            add_paragraph(doc, f'共{total_memos}条备忘录')
            for _, row in memos.head(20).iterrows():
                p = doc.add_paragraph()
                run = p.add_run(f'[{row.get("类型", "—")}] {str(row.get("内容", ""))[:100]}...')
                set_run_font(run, size=10)

    # ========== 保存 ==========
    if output_path:
        doc.save(output_path)
        return output_path
    else:
        return doc


# ==================== Excel报告 ====================

def generate_excel_report(results_dict, output_path=None):
    """
    生成Excel多Sheet质性分析报告

    Sheets:
    1. 概览对比
    2. 维度统计
    3. 情感摘要
    4. 关键事件
    5. 编码片段（各分组原始数据）
    """
    buffer = pd.ExcelWriter(output_path or '/tmp/qda_analysis_report.xlsx',
                           engine='openpyxl')

    # Sheet 1: 概览
    rows = []
    for group_name, data in results_dict.items():
        stats = data.get('stats', {})
        rows.append({
            '分组': group_name,
            '有效记录': stats.get('total_valid_record', 0),
            '原始文档数': stats.get('total_raw', 0),
            '平均文本长度': round(stats.get('text_length_mean', 0)) if stats.get('text_length_mean') else 'N/A',
            '时间跨度': stats.get('time_span', '—'),
        })
    pd.DataFrame(rows).to_excel(buffer, sheet_name='概览对比', index=False)

    # Sheet 2: 维度统计
    dim_rows = []
    for group_name, data in results_dict.items():
        dim_stats = data.get('dimension_stats', {})
        for dim, st in dim_stats.items():
            dim_rows.append({
                '分组': group_name,
                '编码维度': dim,
                '正面提及': st.get('pos', 0),
                '负面提及': st.get('neg', 0),
                '净情感值': st.get('net', 0),
                '正负比': st.get('ratio', 0),
                '总计': st.get('total', 0),
            })
    if dim_rows:
        pd.DataFrame(dim_rows).to_excel(buffer, sheet_name='维度统计', index=False)

    # Sheet 3: 情感摘要
    sent_rows = []
    for group_name, data in results_dict.items():
        s = data.get('sentiment_summary', {})
        sent_rows.append({
            '分组': group_name,
            '强正面率': f"{s.get('intense_positive_rate', 0):.1f}%",
            '强负面率': f"{s.get('intense_negative_rate', 0):.1f}%",
            '转折词率': f"{s.get('transition_rate', 0):.1f}%",
            '隐性负面率': f"{s.get('hidden_dissatisfaction_rate', 0):.1f}%",
        })
    pd.DataFrame(sent_rows).to_excel(buffer, sheet_name='情感摘要', index=False)

    # Sheet 4: 关键事件
    event_rows = []
    for group_name, data in results_dict.items():
        ev = data.get('critical_events', {})
        for event, count in ev.get('positive_events', {}).items():
            event_rows.append({'分组': group_name, '类型': '正向事件', '事件描述': event, '频次': count})
        for event, count in ev.get('negative_events', {}).items():
            event_rows.append({'分组': group_name, '类型': '负向事件', '事件描述': event, '频次': count})
    if event_rows:
        pd.DataFrame(event_rows).to_excel(buffer, sheet_name='关键事件', index=False)

    # Sheet 5: 各分组编码片段
    for group_name, data in results_dict.items():
        df = data.get('df')
        if df is not None:
            sheet = group_name[:28]  # Excel sheet name limit
            # 选择可用列：优先展示文本内容和编码列
            text_cols = [c for c in ['文本内容', 'content', 'text', '文档内容', '原始文本']
                        if c in df.columns]
            cols = text_cols + [c for c in df.columns
                                if c.startswith('代码_') and '_实例' not in c]
            if not cols:
                cols = df.columns.tolist()[:20]
            df[cols].head(200).to_excel(buffer, sheet_name=sheet, index=False)

    buffer.close()
    return output_path or '/tmp/qda_analysis_report.xlsx'


# ==================== Markdown报告 ====================

def generate_markdown_report(results_dict, output_path=None):
    """
    生成Markdown格式质性分析报告（可导入Notion/Obsidian等）
    """
    lines = ['# 质性分析研究报告\n']
    lines.append(f'_生成时间: {datetime.now().strftime("%Y-%m-%d %H:%M")}_\n\n')
    lines.append('---\n\n')

    # 数据概览
    lines.append('## 数据概览\n\n')
    lines.append('| 分组 | 有效记录 | 原始文档数 | 平均文本长度 |\n')
    lines.append('|------|---------|------------|------------|\n')
    for group_name, data in results_dict.items():
        s = data.get('stats', {})
        text_len = round(s.get('text_length_mean', 0)) if s.get('text_length_mean') else '—'
        lines.append(f"| {group_name} | {s.get('total_valid_record', 0)} | "
                    f"{s.get('total_raw', 0)} | {text_len} |\n")

    # 维度分析
    lines.append('\n## 编码维度统计\n\n')
    for group_name, data in results_dict.items():
        dim_stats = data.get('dimension_stats', {})
        if not dim_stats:
            continue
        lines.append(f'### {group_name}\n\n')
        sorted_dims = sorted(dim_stats.items(), key=lambda x: -x[1].get('total', 0))
        lines.append('| 编码维度 | 正面提及 | 负面提及 | 净情感值 | 正负比 |\n')
        lines.append('|----------|---------|---------|---------|-------|\n')
        for dim, st in sorted_dims:
            ratio = f"{st.get('ratio', 0)}:1" if st.get('ratio', 0) > 0 else '—'
            lines.append(f"| {dim} | {st.get('pos', 0)} | {st.get('neg', 0)} | "
                        f"{st.get('net', 0)} | {ratio} |\n")

    # 情感
    lines.append('\n## 情感分析\n\n')
    for group_name, data in results_dict.items():
        s = data.get('sentiment_summary', {})
        lines.append(f'- **{group_name}**: 强正面 {s.get("intense_positive_rate",0):.1f}% | '
                    f'强负面 {s.get("intense_negative_rate",0):.1f}% | '
                    f'隐性负面 {s.get("hidden_dissatisfaction_rate",0):.1f}%\n')

    # 关键事件
    lines.append('\n## 关键事件\n\n')
    for group_name, data in results_dict.items():
        ev = data.get('critical_events', {})
        lines.append(f'### {group_name}\n\n')
        lines.append('**正向触发事件：**\n')
        for event, count in sorted(ev.get('positive_events', {}).items(), key=lambda x: -x[1])[:5]:
            lines.append(f'- {event}: {count}次\n')
        lines.append('\n**负向触发事件：**\n')
        for event, count in sorted(ev.get('negative_events', {}).items(), key=lambda x: -x[1])[:5]:
            lines.append(f'- {event}: {count}次\n')

    content = ''.join(lines)
    if output_path:
        Path(output_path).write_text(content, encoding='utf-8')
    return content
